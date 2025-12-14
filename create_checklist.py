"""
Create print checklist as Word document
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

output_folder = r"c:\Users\samij\Desktop\Heba\Case T 4438-25\PRINT_READY"

def create_checklist():
    """Create print checklist as Word document"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_heading('CHECKLISTA FÖR UTSKRIFT OCH INLÄMNING', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Mål T 4438-25 - Eskilstuna Tingsrätt').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Documents to print
    doc.add_heading('DOKUMENT ATT SKRIVA UT OCH SKRIVA UNDER:', level=1)
    
    # Table
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    
    # Header
    headers = ['Nr', 'Dokument', 'Vem skriver under?', 'Klar?']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # Data
    docs = [
        ('1', '01_SVAROMAL_Mohammad_Sami_Alsharef.docx', 'Mohammad', '☐'),
        ('2', '02_SVAROMAL_Joumana_Alnablsi.docx', 'Joumana', '☐'),
        ('3', '03_FULLMAKT_Joumana_till_Mohammad.docx', 'Joumana + Mohammad', '☐'),
        ('4', '04_BEVISUPPGIFT_WhatsApp.docx', 'Ingen (bilaga)', '☐'),
        ('5', '05_NYCKELBEVIS_Hebas_Skulderkannande.docx', 'Mohammad (intyg)', '☐'),
        ('6', '06_BILAGA_WhatsApp_Full_Chat.docx', 'Ingen (bilaga)', '☐'),
    ]
    
    for i, (nr, doc_name, signer, done) in enumerate(docs, 1):
        table.rows[i].cells[0].text = nr
        table.rows[i].cells[1].text = doc_name
        table.rows[i].cells[2].text = signer
        table.rows[i].cells[3].text = done
    
    doc.add_paragraph()
    
    # Important
    doc.add_heading('VIKTIGT:', level=1)
    
    doc.add_paragraph('Fyll i datum på alla dokument innan underskrift', style='List Bullet')
    doc.add_paragraph('Fullmakten (dokument 3) måste skrivas under av BÅDA Joumana och Mohammad', style='List Bullet')
    doc.add_paragraph('Skriv målnummer T 4438-25 på alla dokument', style='List Bullet')
    doc.add_paragraph('Behåll kopior av allt som skickas', style='List Bullet')
    
    doc.add_paragraph()
    
    # Send to
    doc.add_heading('SKICKA TILL:', level=1)
    
    doc.add_paragraph('Eskilstuna Tingsrätt').runs[0].bold = True
    doc.add_paragraph('Box 333')
    doc.add_paragraph('631 05 Eskilstuna')
    
    doc.add_paragraph()
    doc.add_paragraph('Eller via e-post:').runs[0].bold = True
    doc.add_paragraph('eskilstuna.tingsratt@dom.se')
    
    doc.add_paragraph()
    doc.add_paragraph('Ange:').runs[0].bold = True
    doc.add_paragraph('Målnummer: T 4438-25')
    doc.add_paragraph('Rotel: 2:04')
    
    doc.add_paragraph()
    
    # Deadline
    doc.add_heading('DEADLINE:', level=1)
    
    p = doc.add_paragraph()
    p.add_run('14 dagar från när ni fick stämningen!').bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Datum stämningen mottogs: _______________')
    doc.add_paragraph('Sista datum för svar: _______________')
    
    doc.add_paragraph()
    
    # Final checklist
    doc.add_heading('SLUTLIG CHECKLISTA:', level=1)
    
    doc.add_paragraph('☐ Alla dokument utskrivna')
    doc.add_paragraph('☐ Mohammad har skrivit under sitt svaromål')
    doc.add_paragraph('☐ Joumana har skrivit under sitt svaromål')
    doc.add_paragraph('☐ Joumana har skrivit under fullmakten')
    doc.add_paragraph('☐ Mohammad har skrivit under fullmakten')
    doc.add_paragraph('☐ Datum ifyllt på alla dokument')
    doc.add_paragraph('☐ Kopior sparade hemma')
    doc.add_paragraph('☐ Skickat till domstolen')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Contact
    doc.add_heading('VID FRÅGOR:', level=1)
    
    doc.add_paragraph('Eskilstuna Tingsrätt: 016-15 65 00')
    doc.add_paragraph('Fråga efter: Rotel 2:04')
    
    # Save
    filepath = os.path.join(output_folder, '00_CHECKLISTA_Utskrift.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath

if __name__ == "__main__":
    create_checklist()

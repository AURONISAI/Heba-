"""
Create Word document with full WhatsApp chat export
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

output_folder = r"c:\Users\samij\Desktop\Heba\Case T 4438-25\PRINT_READY"
chat_file = r"c:\Users\samij\Desktop\Heba\Heba chat\_chat.txt"

def create_full_chat_export():
    """Create full chat as Word document"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Title
    title = doc.add_heading('BILAGA: WHATSAPP-KONVERSATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Mål T 4438-25').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Info
    doc.add_paragraph('Mellan:').runs[0].bold = True
    doc.add_paragraph('Mohammad Sami Alsharef ("SAMI\'s JACKET\'s" / "ابو سمير")')
    doc.add_paragraph('och')
    doc.add_paragraph('Heba Alhussien ("اهبة")')
    
    doc.add_paragraph()
    doc.add_paragraph('Period: September 2022 - Februari 2025')
    doc.add_paragraph('Källa: WhatsApp-export')
    
    doc.add_paragraph()
    doc.add_paragraph('=' * 60)
    doc.add_paragraph()
    
    # Key messages to highlight
    key_dates = [
        '2024-12-29',  # Heba admits owing 35,000 kr
        '2024-10-23',  # Heba surprised at 60,000 kr debt
        '2024-09-17',  # "Your Bank Norwegian"
        '2023-01-10',  # Repayment confirmed
        '2025-02-17',  # Warning
        '2025-02-22',  # Blocking
    ]
    
    # Read and add chat
    with open(chat_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    current_date = ""
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check for date change
        date_match = re.match(r'\[(\d{4}-\d{2}-\d{2})', line)
        if date_match:
            msg_date = date_match.group(1)
            if msg_date != current_date:
                current_date = msg_date
                doc.add_paragraph()
                p = doc.add_paragraph(f'--- {current_date} ---')
                p.runs[0].bold = True
                if current_date in key_dates:
                    p.runs[0].font.size = Pt(14)
                doc.add_paragraph()
        
        # Check if this is a key message
        is_key = any(date in line for date in key_dates)
        
        # Add the message
        p = doc.add_paragraph(line)
        if is_key:
            for run in p.runs:
                run.bold = True
        
        # Special highlight for the smoking gun
        if '35 الف كرون' in line or 'سداد دين لابو سمير' in line:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(12)
    
    # Save
    filepath = os.path.join(output_folder, '06_BILAGA_WhatsApp_Full_Chat.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath

if __name__ == "__main__":
    print("Creating full chat export...")
    create_full_chat_export()
    print("Done!")

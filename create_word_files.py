"""
Create Word documents for court submission
Case T 4438-25 - Eskilstuna Tingsrätt
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# Output folder
output_folder = r"c:\Users\samij\Desktop\Heba\Case T 4438-25\PRINT_READY"
os.makedirs(output_folder, exist_ok=True)

def create_svaromal_mohammad():
    """Create Mohammad's Svaromål (court response) as Word document"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_heading('SVAROMÅL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Mål nr T 4438-25')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    # To court
    doc.add_paragraph('Till:')
    p = doc.add_paragraph('Eskilstuna Tingsrätt\nRotel 2:04\n631 86 Eskilstuna')
    
    doc.add_paragraph()
    doc.add_paragraph('Datum: 2025-12-13')
    
    # PARTIES
    doc.add_heading('PARTER', level=1)
    
    doc.add_paragraph('Heba Alhussien (Kärande):').runs[0].bold = True
    doc.add_paragraph('Personnummer: 19890412-1244\nOmbud: Advokat Mayssam Baaken\nAdvokat Mayssam Baaken AB\nNorra Hamngatan 4\n411 14 Göteborg\nE-post: Baaken.law@gmail.com')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Mohammad Sami Alsharef (Svarande 1):').runs[0].bold = True
    doc.add_paragraph('Personnummer: 19860217-5070\nNäshultagatan 12\n632 29 Eskilstuna\nE-post: info@samisjackets.com\nTelefon: 0720147440')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Joumana Alnablsi (Svarande 2):').runs[0].bold = True
    doc.add_paragraph('Personnummer: 19880102-5084\nNäshultagatan 12\n632 29 Eskilstuna\n(Företrädd av Mohammad Sami Alsharef enligt fullmakt)')
    
    # YRKANDEN
    doc.add_heading('YRKANDEN', level=1)
    
    doc.add_paragraph('Mohammad och Joumana yrkar att:')
    doc.add_paragraph('1. Hebas stämning ogillas i sin helhet', style='List Number')
    doc.add_paragraph('2. Heba förpliktas att ersätta Mohammad och Joumanas rättegångskostnader', style='List Number')
    doc.add_paragraph('3. I andra hand: Om rätten skulle finna att någon betalningsskyldighet föreligger, yrkas att beloppet jämkas väsentligt', style='List Number')
    
    # INSTÄLLNING
    doc.add_heading('INSTÄLLNING', level=1)
    
    p = doc.add_paragraph('Mohammad och Joumana ')
    p.add_run('bestrider').bold = True
    p.add_run(' Hebas krav i sin helhet.')
    
    doc.add_paragraph()
    doc.add_paragraph('Bestridda belopp:').runs[0].bold = True
    
    # Table for disputed amounts
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = 'Heba påstår'
    headers[1].text = 'Mohammad & Joumanas svar'
    
    data = [
        ('240 000 kr (kontantlån till Mohammad)', 'BESTRIDS - FALSKT'),
        ('30 000 kr (Swish till Joumana)', 'BESTRIDS - ÅTERBETALAT'),
        ('107 875 kr (Bank Norwegian)', 'BESTRIDS - HEBAS EGET KONTO'),
        ('100 000 kr (Marginalen Bank)', 'BESTRIDS - AFFÄRSVERKSAMHET'),
        ('60 723 kr (ränta)', 'BESTRIDS - INGEN SKULD FINNS'),
        ('Rättegångskostnader', 'BESTRIDS'),
    ]
    
    for i, (col1, col2) in enumerate(data, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
    
    # GRUNDER
    doc.add_heading('GRUNDER FÖR BESTRIDANDET', level=1)
    
    # Ground 1
    doc.add_heading('1. Inget lån har förekommit', level=2)
    doc.add_paragraph('De pengar som Heba påstår sig ha lånat ut var inte lån. Det var del av en informell penningöverföringsverksamhet (s.k. hawala) där Heba fungerade som mellanhand för penningöverföringar mellan Sverige och Syrien.')
    doc.add_paragraph('I denna verksamhet:')
    doc.add_paragraph('Betalade Heba Mohammads svenska räkningar', style='List Bullet')
    doc.add_paragraph('Betalade Mohammad motsvarande belopp i Damaskus till Hebas familj', style='List Bullet')
    doc.add_paragraph('Samma pengar cirkulerade fram och tillbaka flera gånger', style='List Bullet')
    doc.add_paragraph('De 300 000 kr som Heba påstår sig ha gett ut är samma pengar som cirkulerat 3-4 gånger, inte separata lån.')
    
    # Ground 2
    doc.add_heading('2. Pengarna kom från företag, inte Heba personligen', level=2)
    doc.add_paragraph('Enligt Mohammads vetskap kom pengarna på Hebas bankkonto från företag som använde hennes tjänster för penningöverföringar. Dessa företag kan vittna om att pengarna tillhörde Mohammad, inte Heba.')
    
    # Ground 3
    doc.add_heading('3. Heba är skyldig Mohammad och Joumana pengar', level=2)
    doc.add_paragraph('Hebas egna WhatsApp-meddelanden visar att hon var skyldig Mohammad pengar:')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('29 december 2024: ').bold = True
    p.add_run('Heba skrev: "عليي سداد دين لابو سمير 35 الف كرون"')
    doc.add_paragraph('Översättning: "Jag är skyldig Abu Sami [Mohammad] 35 000 kr"')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('23 oktober 2024: ').bold = True
    p.add_run('Heba svarade förvånat på 60 000 kr skuld till Mohammad')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Juni 2023: ').bold = True
    p.add_run('Hebas bokföring visade "حساب ابو سمير ٨٨٥٢ دولار" = "$8 852" skuld till Mohammad')
    
    # Ground 4
    doc.add_heading('4. 30 000 kr till Joumana återbetalades', level=2)
    doc.add_paragraph('Joumana fick 30 000 kr via Swish som del av familjens penningöverföringsverksamhet. Detta belopp återbetalades via Swish, vilket bekräftas i WhatsApp-konversation från januari 2023 där det framgår att "allt som togs via Swish har återbetalats."')
    
    # Ground 5
    doc.add_heading('5. Bank Norwegian var Hebas eget konto', level=2)
    doc.add_paragraph('De 107 875 kr som påstås komma från Bank Norwegian var från Hebas eget bankkonto. I WhatsApp-konversationer kallas det konsekvent "البنك النرويجي تبعك" = "DITT Bank Norwegian". Mohammad och Joumana har aldrig tagit något lån från Bank Norwegian.')
    
    # Ground 6
    doc.add_heading('6. Handskrivna papperet gällde kryptovalutahandel', level=2)
    doc.add_paragraph('Det handskrivna papperet som Heba åberopar var relaterat till kryptovalutahandel, inte ett låneavtal. Heba var involverad i att köpa och sälja kryptovaluta, och papperet var affärsbokföring för dessa transaktioner.')
    
    # Ground 7
    doc.add_heading('7. Hebas trovärdighet är allvarligt skadad', level=2)
    
    doc.add_paragraph('a) Falsk våldtäktsanmälan').runs[0].bold = True
    doc.add_paragraph('Heba har tidigare gjort en polisanmälan om våldtäkt mot Mohammad. Hon påstod att han hade våldtagit henne flera gånger under cirka 2 års tid.')
    p = doc.add_paragraph()
    p.add_run('Polisen lade ned utredningen inom cirka 10 dagar på grund av total brist på bevis.').bold = True
    doc.add_paragraph('Detta visar:')
    doc.add_paragraph('Ett mönster av falska anklagelser', style='List Bullet')
    doc.add_paragraph('Vilja att göra allvarliga osanna påståenden', style='List Bullet')
    doc.add_paragraph('Hämndmotiv', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('b) Hämndmotiv').runs[0].bold = True
    doc.add_paragraph('Tidslinje:')
    
    # Timeline table
    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'
    table2.rows[0].cells[0].text = 'Datum'
    table2.rows[0].cells[1].text = 'Händelse'
    table2.rows[1].cells[0].text = '29 dec 2024'
    table2.rows[1].cells[1].text = 'Heba skriver att hon är skyldig Mohammad 35 000 kr'
    table2.rows[2].cells[0].text = '17 feb 2025'
    table2.rows[2].cells[1].text = 'Mohammad varnar Heba om hennes beteende'
    table2.rows[3].cells[0].text = '22 feb 2025'
    table2.rows[3].cells[1].text = 'Mohammad blockerar Heba på WhatsApp'
    table2.rows[4].cells[0].text = 'Efter 22 feb'
    table2.rows[4].cells[1].text = 'Denna stämning lämnas in'
    
    doc.add_paragraph()
    doc.add_paragraph('Stämningen lämnades in som hämnd för att Heba blev blockerad efter att Mohammad avvisat hennes romantiska närmanden.')
    
    # BEVISNING
    doc.add_heading('BEVISNING', level=1)
    
    doc.add_paragraph('Skriftlig bevisning:').runs[0].bold = True
    
    table3 = doc.add_table(rows=7, cols=3)
    table3.style = 'Table Grid'
    table3.rows[0].cells[0].text = 'Nr'
    table3.rows[0].cells[1].text = 'Bevis'
    table3.rows[0].cells[2].text = 'Bevisar'
    
    evidence = [
        ('1', 'WhatsApp-konversation 29 dec 2024', 'Heba erkänner skuld på 35 000 kr till Mohammad'),
        ('2', 'WhatsApp-konversation okt 2024', 'Heba skyldig 60 000 kr till Mohammad'),
        ('3', 'WhatsApp bokföring juni 2023', 'Heba skyldig $8 852 till Mohammad'),
        ('4', 'WhatsApp jan 2023', 'Återbetalning av Swish-belopp bekräftas'),
        ('5', 'Polisens beslut om nedläggning', 'Hebas falska våldtäktsanmälan avfärdad'),
        ('6', 'WhatsApp 22 feb 2025', 'Blockering som utlöste Hebas hämnd'),
    ]
    
    for i, (nr, bevis, bevisar) in enumerate(evidence, 1):
        table3.rows[i].cells[0].text = nr
        table3.rows[i].cells[1].text = bevis
        table3.rows[i].cells[2].text = bevisar
    
    doc.add_paragraph()
    doc.add_paragraph('Mohammad och Joumana förbehåller sig rätten att åberopa ytterligare bevisning längre fram i processen.')
    
    # SÄRSKILT OM JOUMANA
    doc.add_heading('SÄRSKILT OM JOUMANA ALNABLSI', level=1)
    
    doc.add_paragraph('Joumana har minimal personlig inblandning i denna sak. De 30 000 kr som hon påstås ha fått:')
    doc.add_paragraph('Var del av familjens hawala-verksamhet', style='List Number')
    doc.add_paragraph('Återbetalades via Swish', style='List Number')
    doc.add_paragraph('Hanterades av Mohammad', style='List Number')
    
    doc.add_paragraph('Joumana har utfärdat fullmakt till Mohammad att företräda henne i detta mål och alla relaterade ärenden.')
    doc.add_paragraph('Vi hemställer att rätten överväger att avvisa kravet mot Joumana då hon inte var part i något låneavtal.')
    
    # GENKÄROMÅL
    doc.add_heading('GENKÄROMÅL (FÖRBEHÅLLS)', level=1)
    
    doc.add_paragraph('Mohammad och Joumana förbehåller sig rätten att framställa genkäromål mot Heba avseende:')
    doc.add_paragraph('Utestående skuld: minst 35 000 kr enligt Hebas eget erkännande', style='List Number')
    doc.add_paragraph('Skadestånd för falsk våldtäktsanmälan', style='List Number')
    doc.add_paragraph('Kostnader och utgifter orsakade av denna rättegång', style='List Number')
    
    # SAMMANFATTNING
    doc.add_heading('SAMMANFATTNING', level=1)
    
    doc.add_paragraph('Inget lån existerar - det var en penningöverföringsverksamhet (hawala)', style='List Number')
    doc.add_paragraph('Heba är skyldig Mohammad pengar - 35 000 kr enligt hennes eget skriftliga erkännande', style='List Number')
    doc.add_paragraph('Hebas trovärdighet är obefintlig - falsk våldtäktsanmälan avfärdad av polis', style='List Number')
    doc.add_paragraph('Stämningen är hämnd - lämnad in direkt efter att Heba blockerades', style='List Number')
    doc.add_paragraph('Joumana bör avvisas från målet - minimal inblandning, fullmakt utfärdad till Mohammad', style='List Number')
    
    # DETERMINATION STATEMENT
    doc.add_heading('SLUTORD', level=1)
    
    doc.add_paragraph('Mohammad och Joumana är fast beslutna att försvara sig mot dessa ogrundade anklagelser. Vi är beredda att genomgå hela rättsprocessen, inklusive eventuella överklaganden till hovrätt och Högsta domstolen om så krävs.')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Denna stämning är orättfärdig och grundlös.').bold = True
    p.add_run(' Den är uppenbarligen ett försök till hämnd från en person som själv är skyldig pengar och som tidigare har gjort falska anklagelser som avfärdats av polisen.')
    doc.add_paragraph()
    doc.add_paragraph('Vi kommer att kämpa för vår rätt och för att sanningen ska komma fram, oavsett hur lång tid det tar.')
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Eskilstuna den ______________ 2025')
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Mohammad Sami Alsharef').bold = True
    doc.add_paragraph('Även ombud för Joumana Alnablsi (enligt bifogad fullmakt)')
    
    doc.add_paragraph()
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('(Underskrift)')
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Bilagor:').runs[0].bold = True
    doc.add_paragraph('1. Fullmakt från Joumana Alnablsi')
    doc.add_paragraph('2. WhatsApp-utdrag (med översättningar)')
    doc.add_paragraph('3. Polisens nedläggningsbeslut (om tillgängligt)')
    
    # Save
    filepath = os.path.join(output_folder, '01_SVAROMAL_Mohammad_Sami_Alsharef.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath

def create_svaromal_joumana():
    """Create Joumana's short Svaromål as Word document"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_heading('SVAROMÅL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Mål nr T 4438-25')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    # To court
    doc.add_paragraph('Till:')
    doc.add_paragraph('Eskilstuna Tingsrätt\nRotel 2:04\n631 86 Eskilstuna')
    
    doc.add_paragraph()
    doc.add_paragraph('Datum: 2025-12-13')
    
    # PARTIES
    doc.add_heading('PARTER', level=1)
    
    doc.add_paragraph('Heba Alhussien (Kärande):').runs[0].bold = True
    doc.add_paragraph('Personnummer: 19890412-1244\nOmbud: Advokat Mayssam Baaken')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Joumana Alnablsi (Svarande):').runs[0].bold = True
    doc.add_paragraph('Personnummer: 19880102-5084\nNäshultagatan 12\n632 29 Eskilstuna\nE-post: jomana.alnablsi@hotmail.com\nTelefon: 0735165708')
    
    # OMBUD
    doc.add_heading('OMBUD', level=1)
    
    p = doc.add_paragraph('Jag har utfärdat ')
    p.add_run('fullmakt').bold = True
    p.add_run(' till ')
    p.add_run('Mohammad Sami Alsharef').bold = True
    p.add_run(' (19860217-5070) att företräda mig i detta mål och alla relaterade ärenden. Se bifogad fullmakt.')
    
    doc.add_paragraph('All kommunikation i detta mål ska ske genom mitt ombud Mohammad Sami Alsharef.')
    
    # INSTÄLLNING
    doc.add_heading('INSTÄLLNING', level=1)
    
    p = doc.add_paragraph('Jag ')
    p.add_run('bestrider').bold = True
    p.add_run(' käromålet i sin helhet.')
    
    # KORT MOTIVERING
    doc.add_heading('KORT MOTIVERING', level=1)
    
    p = doc.add_paragraph('1. ')
    p.add_run('Jag har inte ingått något låneavtal').bold = True
    p.add_run(' med Heba Alhussien.')
    
    doc.add_paragraph()
    p = doc.add_paragraph('2. ')
    p.add_run('De 30 000 kr').bold = True
    p.add_run(' som påstås ha överförts till mig var del av en informell penningöverföringsverksamhet (hawala) som hanterades av Mohammad.')
    
    doc.add_paragraph()
    p = doc.add_paragraph('3. ')
    p.add_run('Beloppet återbetalades').bold = True
    p.add_run(' via Swish-överföringar, vilket framgår av WhatsApp-konversationer.')
    
    doc.add_paragraph()
    p = doc.add_paragraph('4. ')
    p.add_run('Jag hade minimal personlig inblandning').bold = True
    p.add_run(' i de ekonomiska arrangemangen mellan Mohammad och Heba.')
    
    doc.add_paragraph()
    doc.add_paragraph('5. Jag ansluter mig i övrigt till det svaromål som inges av mitt ombud Mohammad Sami Alsharef.')
    
    # HEMSTÄLLAN
    doc.add_heading('HEMSTÄLLAN', level=1)
    
    doc.add_paragraph('Jag hemställer att:')
    p = doc.add_paragraph('1. ')
    p.add_run('Hebas stämning mot mig ogillas').bold = True
    
    p = doc.add_paragraph('2. ')
    p.add_run('Heba förpliktas ersätta mina rättegångskostnader').bold = True
    
    doc.add_paragraph('3. Alternativt att kravet mot mig avvisas då jag inte var part i något låneavtal')
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Eskilstuna den ______________ 2025')
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Joumana Alnablsi').bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('(Underskrift)')
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Bilaga:').runs[0].bold = True
    doc.add_paragraph('Fullmakt till Mohammad Sami Alsharef')
    
    # Save
    filepath = os.path.join(output_folder, '02_SVAROMAL_Joumana_Alnablsi.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath

def create_fullmakt():
    """Create Fullmakt (Power of Attorney) as Word document"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_heading('FULLMAKT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Rättegångsfullmakt')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    # FULLMAKTSGIVARE
    doc.add_heading('FULLMAKTSGIVARE', level=1)
    
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Table Grid'
    table1.rows[0].cells[0].text = 'Namn:'
    table1.rows[0].cells[1].text = 'Joumana Alnablsi'
    table1.rows[1].cells[0].text = 'Personnummer:'
    table1.rows[1].cells[1].text = '19880102-5084'
    table1.rows[2].cells[0].text = 'Adress:'
    table1.rows[2].cells[1].text = 'Näshultagatan 12, 632 29 Eskilstuna'
    table1.rows[3].cells[0].text = 'Telefon:'
    table1.rows[3].cells[1].text = '0735165708'
    table1.rows[4].cells[0].text = 'E-post:'
    table1.rows[4].cells[1].text = 'jomana.alnablsi@hotmail.com'
    
    doc.add_paragraph()
    
    # FULLMAKTSTAGARE
    doc.add_heading('FULLMAKTSTAGARE (OMBUD)', level=1)
    
    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Table Grid'
    table2.rows[0].cells[0].text = 'Namn:'
    table2.rows[0].cells[1].text = 'Mohammad Sami Alsharef'
    table2.rows[1].cells[0].text = 'Personnummer:'
    table2.rows[1].cells[1].text = '19860217-5070'
    table2.rows[2].cells[0].text = 'Adress:'
    table2.rows[2].cells[1].text = 'Näshultagatan 12, 632 29 Eskilstuna'
    table2.rows[3].cells[0].text = 'Telefon:'
    table2.rows[3].cells[1].text = '0720147440'
    table2.rows[4].cells[0].text = 'E-post:'
    table2.rows[4].cells[1].text = 'info@samisjackets.com'
    table2.rows[5].cells[0].text = 'Relation:'
    table2.rows[5].cells[1].text = 'Före detta sambo'
    
    doc.add_paragraph()
    
    # FULLMAKTENS OMFATTNING
    doc.add_heading('FULLMAKTENS OMFATTNING', level=1)
    
    p = doc.add_paragraph('Jag, ')
    p.add_run('Joumana Alnablsi').bold = True
    p.add_run(', ger härmed ')
    p.add_run('Mohammad Sami Alsharef').bold = True
    p.add_run(' fullmakt att för min räkning:')
    
    doc.add_paragraph()
    doc.add_paragraph('Rättegångsfullmakt (Mål T 4438-25):').runs[0].bold = True
    
    doc.add_paragraph('Företräda mig som svarande i mål nr T 4438-25 vid Eskilstuna Tingsrätt', style='List Bullet')
    doc.add_paragraph('Avge svaromål, yttranden och övriga inlagor', style='List Bullet')
    doc.add_paragraph('Närvara vid muntlig förberedelse och huvudförhandling', style='List Bullet')
    doc.add_paragraph('Ingå förlikning å mina vägnar', style='List Bullet')
    doc.add_paragraph('Ta emot delgivning av domar och beslut', style='List Bullet')
    doc.add_paragraph('Överklaga domar och beslut', style='List Bullet')
    doc.add_paragraph('Vidta alla åtgärder som krävs för att tillvarata mina intressen i målet', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Generalfullmakt avseende detta ärende:').runs[0].bold = True
    
    doc.add_paragraph('Företräda mig inför alla myndigheter i anledning av detta mål', style='List Bullet')
    doc.add_paragraph('Ta emot och kvittera handlingar', style='List Bullet')
    doc.add_paragraph('Kommunicera med motparten och dennes ombud', style='List Bullet')
    doc.add_paragraph('Fatta beslut om processåtgärder', style='List Bullet')
    
    # SÄRSKILDA VILLKOR
    doc.add_heading('SÄRSKILDA VILLKOR', level=1)
    
    doc.add_paragraph('1. Denna fullmakt är obegränsad i tid och gäller tills den uttryckligen återkallas av mig.')
    doc.add_paragraph('2. Fullmaktstagaren har rätt att substituera fullmakten till juridiskt ombud om sådant anlitas.')
    doc.add_paragraph('3. Fullmaktstagaren får ingå förlikning å mina vägnar.')
    doc.add_paragraph('4. Jag förbinder mig att godkänna de åtgärder som fullmaktstagaren vidtar inom ramen för denna fullmakt.')
    
    # SKÄL TILL FULLMAKTEN
    doc.add_heading('SKÄL TILL FULLMAKTEN', level=1)
    
    doc.add_paragraph('Jag utfärdar denna fullmakt eftersom:')
    doc.add_paragraph('1. Mohammad Sami Alsharef har fullständig kunskap om de faktiska omständigheterna i målet')
    doc.add_paragraph('2. Jag hade minimal personlig inblandning i de ekonomiska arrangemangen med Heba Alhussien')
    doc.add_paragraph('3. Det är praktiskt att ha en gemensam företrädare för båda svarandena')
    
    # UNDERSKRIFTER
    doc.add_heading('UNDERSKRIFTER', level=1)
    
    doc.add_paragraph('Fullmaktsgivare:').runs[0].bold = True
    doc.add_paragraph()
    doc.add_paragraph('Ort och datum: Eskilstuna den ______________ 2025')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('Joumana Alnablsi (19880102-5084)')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc.add_paragraph('Fullmaktstagare (bekräftelse av mottagande):').runs[0].bold = True
    doc.add_paragraph()
    doc.add_paragraph('Jag, Mohammad Sami Alsharef, bekräftar att jag har mottagit denna fullmakt och åtar mig att företräda Joumana Alnablsi i enlighet med dess villkor.')
    doc.add_paragraph()
    doc.add_paragraph('Ort och datum: Eskilstuna den ______________ 2025')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('Mohammad Sami Alsharef (19860217-5070)')
    
    # Save
    filepath = os.path.join(output_folder, '03_FULLMAKT_Joumana_till_Mohammad.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath

def create_whatsapp_evidence():
    """Create WhatsApp evidence document as Word file"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_heading('BEVISUPPGIFT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('WhatsApp-konversationer')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    
    subtitle2 = doc.add_paragraph('Mål T 4438-25')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Document info
    doc.add_heading('DOKUMENTINFORMATION', level=1)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Källa:'
    table.rows[0].cells[1].text = 'WhatsApp-konversation mellan Mohammad och Heba'
    table.rows[1].cells[0].text = 'Period:'
    table.rows[1].cells[1].text = '2022 - februari 2025'
    table.rows[2].cells[0].text = 'Språk:'
    table.rows[2].cells[1].text = 'Arabiska (med svensk översättning)'
    table.rows[3].cells[0].text = 'Relevans:'
    table.rows[3].cells[1].text = 'Bevisar att Heba var skyldig Mohammad pengar, inte tvärtom'
    
    # NYCKELBEVIS
    doc.add_heading('NYCKELBEVIS', level=1)
    
    # BEVIS 1
    doc.add_heading('BEVIS 1: Hebas skulderkännande (29 december 2024)', level=2)
    
    doc.add_paragraph('Datum: 2024-12-29, 17:25:28').runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Original (arabiska):').runs[0].italic = True
    p = doc.add_paragraph()
    p.add_run('عليي سداد دين لابو سمير 35 الف كرون . لتاريخ اليوم بشهر ديسمبر وتم دفع الرينتا من ابو سمير كاش 1800 كرون ل هبه').italic = True
    
    doc.add_paragraph()
    doc.add_paragraph('Svensk översättning:').runs[0].italic = True
    p = doc.add_paragraph()
    p.add_run('"Jag är skyldig Abu Sami [Mohammad] 35 000 kronor. Per dagens datum i december, och hyran betalades av Abu Sami kontant 1 800 kronor till Heba."').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph('Bevisvärde: ')
    p.add_run('AVGÖRANDE').bold = True
    p.add_run(' - Heba erkänner skriftligen att hon är skyldig Mohammad pengar endast cirka 2 månader innan hon lämnar in stämningen.')
    
    # BEVIS 2
    doc.add_heading('BEVIS 2: Heba skyldig 60 000 kr (23 oktober 2024)', level=2)
    
    doc.add_paragraph('Datum: 2024-10-23, 11:58:41 - 11:59:20').runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Original (arabiska):').runs[0].italic = True
    p = doc.add_paragraph()
    p.add_run('كيف عليي ٦٠ الف\n🙄🙄🙄\nشوووو الك عليي ٦٠ الف ؟\nكيف').italic = True
    
    doc.add_paragraph()
    doc.add_paragraph('Svensk översättning:').runs[0].italic = True
    p = doc.add_paragraph()
    p.add_run('"Hur är jag skyldig 60 000? 🙄🙄🙄 VADÅ du säger jag är skyldig dig 60 000? Hur?"').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph('Bevisvärde: ')
    p.add_run('STARKT').bold = True
    p.add_run(' - Heba ifrågasätter beloppet, inte faktumet att hon är skyldig pengar. Detta bekräftar att Mohammad förde bokföring över hennes skulder till honom.')
    
    # BEVIS 3
    doc.add_heading('BEVIS 3: "DITT Bank Norwegian" (17 september 2024)', level=2)
    
    doc.add_paragraph('Datum: 2024-09-17, 09:22:57').runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Original (arabiska):').runs[0].italic = True
    p = doc.add_paragraph()
    p.add_run('فوتي حوليلي من النوريجين تبعك فورا 35000 kr لحسابي').italic = True
    
    doc.add_paragraph()
    doc.add_paragraph('Svensk översättning:').runs[0].italic = True
    p = doc.add_paragraph()
    p.add_run('"Gå in och överför till mig från DITT Norwegian [Bank Norwegian] omedelbart 35 000 kr till mitt konto"').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph('Bevisvärde: ')
    p.add_run('STARKT').bold = True
    p.add_run(' - Visar att Bank Norwegian var Hebas eget konto ("تبعك" = "ditt"), inte ett lån hon tagit för Mohammad och Joumana.')
    
    # BEVIS 4
    doc.add_heading('BEVIS 4: Återbetalning bekräftad (januari 2023)', level=2)
    
    doc.add_paragraph('Datum: 2023-01-10, 14:16:02').runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Original (arabiska):').runs[0].italic = True
    p = doc.add_paragraph()
    p.add_run('ربي يسلم ايديكي مبدائيا رجعتلك كل شي اخذتو منك سويش').italic = True
    
    doc.add_paragraph()
    doc.add_paragraph('Svensk översättning:').runs[0].italic = True
    p = doc.add_paragraph()
    p.add_run('"Gud välsigne dina händer, jag har i princip lämnat tillbaka allt jag tog från dig via Swish"').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph('Bevisvärde: ')
    p.add_run('STÖDJANDE').bold = True
    p.add_run(' - Visar att pengar som togs via Swish (inklusive de 30 000 kr till Joumana) återbetalades.')
    
    # BEVIS 5 - Blockering
    doc.add_heading('BEVIS 5: Blockeringen och hämnden (februari 2025)', level=2)
    
    doc.add_paragraph('Datum: 2025-02-17, 11:15:08').runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Original (arabiska):').runs[0].italic = True
    p = doc.add_paragraph()
    p.add_run('أنا نبهتك بالحكي بالهدوء أنت شكلك ما فهمت بالحكي بالهدوء هلا شوفي الوش الثاني بقى').italic = True
    
    doc.add_paragraph()
    doc.add_paragraph('Svensk översättning:').runs[0].italic = True
    p = doc.add_paragraph()
    p.add_run('"Jag varnade dig med lugnt tal, du verkade inte förstå med lugnt tal, nu får du se den andra sidan"').bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Datum: 2025-02-22, 09:43:18').runs[0].bold = True
    doc.add_paragraph('Heba: "أنا بشو غلطت معك" = "Vad har jag gjort för fel mot dig?"')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Därefter: "You blocked this contact" (Du blockerade denna kontakt)').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph('Bevisvärde: ')
    p.add_run('AVGÖRANDE FÖR HÄMNDMOTIV').bold = True
    p.add_run(' - Visar att stämningen lämnades in som hämnd efter blockering.')
    
    # SLUTSATS
    doc.add_heading('SLUTSATS', level=1)
    
    doc.add_paragraph('WhatsApp-konversationerna visar tydligt att:')
    doc.add_paragraph('Heba erkände skriftligen att hon var skyldig Mohammad 35 000 kr (29 dec 2024)', style='List Number')
    doc.add_paragraph('Bank Norwegian var Hebas eget konto, inte ett lån till Mohammad', style='List Number')
    doc.add_paragraph('Pengar som togs via Swish återbetalades', style='List Number')
    doc.add_paragraph('Stämningen lämnades in som hämnd efter att Mohammad blockerade Heba', style='List Number')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Hebas påståenden i stämningen motsägs direkt av hennes egna skriftliga uttalanden.').bold = True
    
    # Save
    filepath = os.path.join(output_folder, '04_BEVISUPPGIFT_WhatsApp.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath

def create_key_chat_excerpt():
    """Create a document with the key chat message - the smoking gun"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_heading('UTDRAG FRÅN WHATSAPP', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Hebas skulderkännande')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    
    doc.add_paragraph()
    
    # Info box
    p = doc.add_paragraph('Mål: T 4438-25')
    p = doc.add_paragraph('Datum för meddelande: 29 december 2024, kl 17:25:28')
    p = doc.add_paragraph('Avsändare: Heba Alhussien (اهبة)')
    p = doc.add_paragraph('Mottagare: Mohammad Sami Alsharef (SAMI\'s JACKET\'s)')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # The message
    doc.add_heading('ORIGINALMEDDELANDE (ARABISKA):', level=1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('عليي سداد دين لابو سمير 35 الف كرون . لتاريخ اليوم بشهر ديسمبر وتم دفع الرينتا من ابو سمير كاش 1800 كرون ل هبه')
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc.add_heading('SVENSK ÖVERSÄTTNING:', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('"Jag är skyldig Abu Sami [Mohammad] 35 000 kronor. Per dagens datum i december, och hyran betalades av Abu Sami kontant 1 800 kronor till Heba."')
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Significance
    doc.add_heading('BETYDELSE FÖR MÅLET:', level=1)
    
    doc.add_paragraph('Detta meddelande skrevs av Heba Alhussien endast cirka 2 månader innan hon lämnade in sin stämningsansökan mot Mohammad och Joumana.')
    doc.add_paragraph()
    doc.add_paragraph('I meddelandet erkänner Heba uttryckligen att:')
    doc.add_paragraph('Hon är skyldig Mohammad ("Abu Sami") pengar', style='List Number')
    doc.add_paragraph('Skulden uppgår till 35 000 kr', style='List Number')
    doc.add_paragraph('Mohammad betalade hennes hyra (1 800 kr kontant)', style='List Number')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Detta motsäger helt Hebas påstående i stämningen att Mohammad är skyldig henne pengar.').bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Certification
    doc.add_paragraph('Jag intygar att ovanstående är en korrekt återgivning av WhatsApp-meddelandet:')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Eskilstuna den ______________ 2025')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('Mohammad Sami Alsharef')
    
    # Save
    filepath = os.path.join(output_folder, '05_NYCKELBEVIS_Hebas_Skulderkannande.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath

# Run all
if __name__ == "__main__":
    print("=" * 60)
    print("CREATING WORD DOCUMENTS FOR COURT SUBMISSION")
    print("Case T 4438-25 - Eskilstuna Tingsrätt")
    print("=" * 60)
    print()
    
    create_svaromal_mohammad()
    create_svaromal_joumana()
    create_fullmakt()
    create_whatsapp_evidence()
    create_key_chat_excerpt()
    
    print()
    print("=" * 60)
    print(f"✅ ALL DOCUMENTS CREATED IN: {output_folder}")
    print("=" * 60)
    print()
    print("DOCUMENTS READY FOR PRINTING:")
    print("1. 01_SVAROMAL_Mohammad_Sami_Alsharef.docx - Mohammad signs")
    print("2. 02_SVAROMAL_Joumana_Alnablsi.docx - Joumana signs")
    print("3. 03_FULLMAKT_Joumana_till_Mohammad.docx - BOTH sign")
    print("4. 04_BEVISUPPGIFT_WhatsApp.docx - Evidence (no signature)")
    print("5. 05_NYCKELBEVIS_Hebas_Skulderkannande.docx - Key evidence")

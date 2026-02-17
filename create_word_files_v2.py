"""
Create Word documents for court submission - UPDATED AFTER FIRST HEARING
Case T 4438-25 - Eskilstuna Tingsrätt
Updated: 2026-02-16 (after first court hearing)
Next hearing: 2026-05-18
Evidence deadline: 2026-03-20
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Output folder
output_folder = r"c:\Users\samij\Desktop\Heba\Case T 4438-25\PRINT_READY"
os.makedirs(output_folder, exist_ok=True)


def set_arabic_font(run, font_name='Arial', size=12):
    """Set font that supports Arabic text"""
    run.font.name = font_name
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_arabic_with_translation(doc, arabic_text, swedish_text, context=""):
    """Add Arabic text with Swedish translation in a clear format"""
    if context:
        doc.add_paragraph(context)
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    table.rows[0].cells[0].text = 'ARABISKA (Original)'
    table.rows[0].cells[1].text = 'SVENSKA (Översättning)'
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    arabic_cell = table.rows[1].cells[0]
    swedish_cell = table.rows[1].cells[1]
    
    arabic_para = arabic_cell.paragraphs[0]
    arabic_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    arabic_run = arabic_para.add_run(arabic_text)
    set_arabic_font(arabic_run, 'Arial', 11)
    
    swedish_para = swedish_cell.paragraphs[0]
    swedish_run = swedish_para.add_run(swedish_text)
    swedish_run.font.name = 'Arial'
    swedish_run.font.size = Pt(11)
    swedish_run.bold = True
    
    doc.add_paragraph()


def setup_doc():
    """Create a Document with standard settings"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    return doc


# =============================================================================
# 1. SVAROMÅL - MOHAMMAD (UPDATED)
# =============================================================================
def create_svaromal_mohammad():
    """Create Mohammad's Svaromål - UPDATED after first hearing 2026-02-16"""
    doc = setup_doc()
    
    # Title
    title = doc.add_heading('SVAROMÅL OCH GENKÄROMÅL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Mål nr T 4438-25')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Uppdaterad inlaga efter muntlig förberedelse 2026-02-16')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    p.runs[0].italic = True
    
    doc.add_paragraph()
    
    # To court
    doc.add_paragraph('Till:')
    doc.add_paragraph('Eskilstuna Tingsrätt\nRotel 2:04\n631 86 Eskilstuna')
    doc.add_paragraph()
    doc.add_paragraph('Datum: 2026-02-16')
    
    # PARTIES
    doc.add_heading('PARTER', level=1)
    
    doc.add_paragraph('Kärande:').runs[0].bold = True
    doc.add_paragraph('Heba Alhussien\nPersonnummer: 19890412-1244\nOmbud: Advokat Mayssam Baaken\nAdvokat Mayssam Baaken AB\nNorra Hamngatan 4, 411 14 Göteborg\nE-post: Baaken.law@gmail.com')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Svarande 1:').runs[0].bold = True
    doc.add_paragraph('Mohammad Sami Alsharef\nPersonnummer: 19860217-5070\nNäshultagatan 12, 632 29 Eskilstuna\nE-post: info@samisjackets.com\nTelefon: 0720147440')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Svarande 2:').runs[0].bold = True
    doc.add_paragraph('Joumana Alnablsi\nPersonnummer: 19880102-5084\nNäshultagatan 12, 632 29 Eskilstuna\n(Företrädd av Mohammad Sami Alsharef enligt fullmakt)')
    
    # PROCESSHISTORIK
    doc.add_heading('PROCESSHISTORIK', level=1)
    
    table_proc = doc.add_table(rows=5, cols=2)
    table_proc.style = 'Table Grid'
    table_proc.rows[0].cells[0].text = 'Stämning mottagen'
    table_proc.rows[0].cells[1].text = '2025 (exakt datum i handlingar)'
    table_proc.rows[1].cells[0].text = 'Svaromål inlämnat'
    table_proc.rows[1].cells[1].text = '2025-12-14'
    table_proc.rows[2].cells[0].text = 'Första muntliga förberedelse'
    table_proc.rows[2].cells[1].text = '2026-02-16'
    table_proc.rows[3].cells[0].text = 'Frist för ny bevisning'
    table_proc.rows[3].cells[1].text = '2026-03-20'
    table_proc.rows[4].cells[0].text = 'Nästa förhandling'
    table_proc.rows[4].cells[1].text = '2026-05-18'
    
    doc.add_paragraph()
    
    # YRKANDEN
    doc.add_heading('YRKANDEN', level=1)
    
    doc.add_paragraph('Mohammad och Joumana yrkar att:')
    doc.add_paragraph('1. Käromålet ogillas i sin helhet', style='List Number')
    doc.add_paragraph('2. Heba förpliktas att ersätta Mohammad och Joumanas rättegångskostnader', style='List Number')
    
    doc.add_paragraph()
    doc.add_heading('Genkäromål (400 000 kr)', level=2)
    doc.add_paragraph('Mohammad yrkar i genkäromål att Heba Alhussien förpliktas att betala sammanlagt 400 000 kr till Mohammad Sami Alsharef, fördelat enligt följande:')
    doc.add_paragraph()
    
    table_gk = doc.add_table(rows=7, cols=3)
    table_gk.style = 'Table Grid'
    table_gk.rows[0].cells[0].text = 'Post'
    table_gk.rows[0].cells[1].text = 'Belopp'
    table_gk.rows[0].cells[2].text = 'Rättslig grund'
    for cell in table_gk.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    
    gk_data = [
        ('A. Hebas skuld (erkännande)', '35 000 kr', 'Ensidigt skulderkännande 2024-12-29'),
        ('B. Kränkningsersättning – falsk våldtäktsanmälan', '150 000 kr', '2 kap 3 § SkL via BrB 15:7 (falsk tillvitelse)'),
        ('C. Skadestånd – rättegångsmissbruk', '100 000 kr', '2 kap 2 § SkL, jfr BrB 15:2 (osann partsutsaga)'),
        ('D. Ekonomisk skada', '75 000 kr', '2 kap 2 § SkL (översättning, förlorad arbetstid, inkomstförlust)'),
        ('E. Sveda och värk (psykiskt lidande)', '40 000 kr', '5 kap 1 § SkL (personskada – psykiskt lidande)'),
        ('TOTALT', '400 000 kr', '+ ränta + rättegångskostnader'),
    ]
    
    for i, (post, belopp, grund) in enumerate(gk_data, 1):
        table_gk.rows[i].cells[0].text = post
        table_gk.rows[i].cells[1].text = belopp
        table_gk.rows[i].cells[2].text = grund
        if i == 6:
            for cell in table_gk.rows[i].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Jämte ränta enligt 6 § räntelagen (1975:635) från respektive skadetillfälle till dess betalning sker.')
    doc.add_paragraph('Heba förpliktas att ersätta samtliga rättegångskostnader i såväl käromålet som genkäromålet.')
    
    # INSTÄLLNING
    doc.add_heading('INSTÄLLNING', level=1)
    
    p = doc.add_paragraph('Mohammad och Joumana ')
    p.add_run('bestrider').bold = True
    p.add_run(' Hebas samtliga krav i sin helhet. ')
    p.add_run('Ingen förlikning accepteras.').bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Bestridda belopp:').runs[0].bold = True
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = 'Post'
    headers[1].text = 'Heba påstår'
    headers[2].text = 'Svarandenas inställning'
    for cell in headers:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    
    data = [
        ('Kontantlån till Mohammad', '240 000 kr', 'BESTRIDS – inget lån, hawala-verksamhet'),
        ('Swish till Joumana', '30 000 kr', 'BESTRIDS – återbetalat via Swish'),
        ('Bank Norwegian', '107 875 kr', 'BESTRIDS – Hebas eget konto'),
        ('Marginalen Bank', '100 000 kr', 'BESTRIDS – affärsverksamhet, ej lån'),
        ('Ränta', '60 723 kr', 'BESTRIDS – ingen skuld finns'),
        ('Rättegångskostnader', '', 'BESTRIDS'),
    ]
    
    for i, (post, belopp, svar) in enumerate(data, 1):
        table.rows[i].cells[0].text = post
        table.rows[i].cells[1].text = belopp
        table.rows[i].cells[2].text = svar
    
    # GRUNDER
    doc.add_heading('GRUNDER FÖR BESTRIDANDET', level=1)
    
    # Ground 1
    doc.add_heading('1. Inga lån – informell penningöverföring (hawala)', level=2)
    doc.add_paragraph('De påstådda lånen var del av en informell penningöverföringsverksamhet (hawala) där Heba fungerade som mellanhand mellan Sverige och Syrien. Heba betalade svenska räkningar (fakturor, Klarna), Mohammad betalade motsvarande i Damaskus i syriska pund. De 300 000+ kr som Heba åberopar är samma pengar som cirkulerat 3–4 gånger – inte separata lån.')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('NY BEVISNING – Abu Mohammad (penningväxlare/صراف):').bold = True
    doc.add_paragraph('WhatsApp-konversation (518 meddelanden, dec 2023–sep 2024) mellan Heba och Abu Mohammad bevisar systematisk hawala-verksamhet:')
    doc.add_paragraph('Dollarväxlingar i Damaskus, belopp i syriska pund (11,5 milj. SYP)', style='List Bullet')
    doc.add_paragraph('USDT/kryptovaluta, Klarna-/Nordea-fakturor, kvitton', style='List Bullet')
    doc.add_paragraph('60+ ljudmeddelanden med transaktionsinstruktioner', style='List Bullet')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Abu Mohammad-chatten bevisar att Hebas verksamhet var professionell penningväxling – inte privata utlåningar.').bold = True
    
    # Ground 2
    doc.add_heading('2. Heba är skyldig Mohammad pengar', level=2)
    doc.add_paragraph('Hebas egna WhatsApp-meddelanden bevisar att hon var skyldig Mohammad pengar:')
    doc.add_paragraph()
    
    # BEVIS B1
    p = doc.add_paragraph()
    p.add_run('BEVIS B1 – 29 december 2024 (Hebas skulderkännande):').bold = True
    
    add_arabic_with_translation(
        doc,
        'عليي سداد دين لابو سمير 35 الف كرون',
        '"Jag är skyldig Abu Sami [Mohammad] 35 000 kr"'
    )
    
    p = doc.add_paragraph()
    p.add_run('→ Skrivet av Heba 2 månader före stämningen. Detta är grunden för genkäromålet.').italic = True
    
    # BEVIS B2
    p = doc.add_paragraph()
    p.add_run('BEVIS B2 – 23 oktober 2024 (Heba skyldig 60 000 kr):').bold = True
    
    add_arabic_with_translation(
        doc,
        'كيف عليي ٦٠ الف ؟ شوووو الك عليي ٦٠ الف',
        '"Hur är jag skyldig 60 000? VADÅ du säger jag är skyldig dig 60 000?"'
    )
    
    p = doc.add_paragraph()
    p.add_run('→ Heba ifrågasätter beloppet men inte att hon är skyldig.').italic = True
    
    # BEVIS B3
    p = doc.add_paragraph()
    p.add_run('BEVIS B3 – Juni 2023 (Hebas bokföring):').bold = True
    
    add_arabic_with_translation(
        doc,
        'حساب ابو سمير ٨٨٥٢ دولار',
        '"Abu Samis [Mohammads] konto: 8 852 dollar" (skuld till Mohammad)'
    )
    
    # Ground 3
    doc.add_heading('3. 30 000 kr till Joumana återbetalades', level=2)
    doc.add_paragraph('Joumana fick 30 000 kr via Swish som del av hawala-verksamheten. Beloppet återbetalades via Swish:')
    
    add_arabic_with_translation(
        doc,
        'ربي يسلم ايديكي مبدائيا رجعتلك كل شي اخذتو منك سويش',
        '"Gud välsigne dina händer, jag har i princip lämnat tillbaka allt jag tog från dig via Swish"'
    )
    
    # Ground 4
    doc.add_heading('4. Bank Norwegian var Hebas eget konto', level=2)
    doc.add_paragraph('De 107 875 kr från Bank Norwegian var från Hebas eget konto:')
    
    add_arabic_with_translation(
        doc,
        'البنك النرويجي تبعك',
        '"DITT Bank Norwegian" (= Hebas eget konto, inte ett lån till Mohammad/Joumana)'
    )
    
    doc.add_paragraph('Mohammad och Joumana har aldrig tagit något lån från Bank Norwegian.')
    
    # Ground 5
    doc.add_heading('5. Handskrivna papperet – kryptovalutahandel', level=2)
    doc.add_paragraph('Det handskrivna papperet var relaterat till USDT-kryptohandel, inte låneavtal. Abu Mohammad-chatten bekräftar att USDT-handel ingår i Hebas hawala-verksamhet.')
    
    # Ground 6
    doc.add_heading('6. Hebas trovärdighet är allvarligt skadad', level=2)
    
    doc.add_paragraph('a) Falsk våldtäktsanmälan').runs[0].bold = True
    doc.add_paragraph('Heba anmälde Mohammad för våldtäkt – påstod upprepade övergrepp under ~2 år. Polisen lade ned utredningen inom ~10 dagar p.g.a. total brist på bevis. Detta visar ett mönster av falska anklagelser och hämndmotiv.')
    
    doc.add_paragraph()
    doc.add_paragraph('b) Hämndmotiv – tidslinje').runs[0].bold = True
    
    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Table Grid'
    table2.rows[0].cells[0].text = 'Datum'
    table2.rows[0].cells[1].text = 'Händelse'
    for cell in table2.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    table2.rows[1].cells[0].text = '29 dec 2024'
    table2.rows[1].cells[1].text = 'Heba skriver att hon är skyldig Mohammad 35 000 kr'
    table2.rows[2].cells[0].text = '17 feb 2025'
    table2.rows[2].cells[1].text = 'Mohammad varnar Heba om hennes beteende'
    table2.rows[3].cells[0].text = 'Februari 2025'
    table2.rows[3].cells[1].text = 'Heba polisanmäler Mohammad för våldtäkt (nedlagd ~10 dagar)'
    table2.rows[4].cells[0].text = 'Mars 2025'
    table2.rows[4].cells[1].text = 'Heba lämnar in stämning – T 4438-25'
    table2.rows[5].cells[0].text = '16 feb 2026'
    table2.rows[5].cells[1].text = 'Första muntliga förberedelse'
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Stämningen är hämnd efter att polisen avfärdade Hebas falska anmälan.').bold = True
    
    # BEVISNING
    doc.add_heading('BEVISNING', level=1)
    
    doc.add_paragraph('OBS: All arabisk bevisning kommer att förses med auktoriserad översättning enligt rättens instruktioner från 2026-02-16.').runs[0].italic = True
    doc.add_paragraph()
    
    doc.add_paragraph('Skriftlig bevisning:').runs[0].bold = True
    
    table3 = doc.add_table(rows=11, cols=3)
    table3.style = 'Table Grid'
    table3.rows[0].cells[0].text = 'Nr'
    table3.rows[0].cells[1].text = 'Bevis'
    table3.rows[0].cells[2].text = 'Bevisar'
    for cell in table3.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    
    evidence = [
        ('B1', 'WhatsApp 29 dec 2024 – Hebas skulderkännande', 'Heba erkänner skuld 35 000 kr till Mohammad'),
        ('B2', 'WhatsApp okt 2024 – "60 000 kr"', 'Heba skyldig 60 000 kr till Mohammad'),
        ('B3', 'WhatsApp juni 2023 – bokföring $8 852', 'Hebas egen bokföring visar skuld till Mohammad'),
        ('B4', 'WhatsApp jan 2023 – återbetalning Swish', '30 000 kr till Joumana återbetalades'),
        ('B5', 'Polisens nedläggningsbeslut', 'Hebas falska våldtäktsanmälan avfärdad'),
        ('B6', 'Abu Mohammad-chatt – dollarväxlingar', 'Heba bedrev professionell hawala via penningväxlare'),
        ('B7', 'Abu Mohammad-chatt – syriska pund', 'Stora belopp (11,5 miljoner SYP) i transaktioner'),
        ('B8', 'Abu Mohammad-chatt – USDT/krypto', 'Kryptovalutahandel som del av hawala-verksamhet'),
        ('B9', 'Abu Mohammad-chatt – faktura-/Klarna-betalningar', 'Systematisk fakturering via Nordea/Klarna'),
        ('B10', 'Abu Mohammad-chatt – ljudfiler och foton', '60+ ljud, 80+ foton med transaktionsbevis'),
    ]
    
    for i, (nr, bevis, bevisar) in enumerate(evidence, 1):
        table3.rows[i].cells[0].text = nr
        table3.rows[i].cells[1].text = bevis
        table3.rows[i].cells[2].text = bevisar
    
    doc.add_paragraph()
    
    doc.add_paragraph('Muntlig bevisning:').runs[0].bold = True
    
    table4 = doc.add_table(rows=3, cols=3)
    table4.style = 'Table Grid'
    table4.rows[0].cells[0].text = 'Nr'
    table4.rows[0].cells[1].text = 'Vittne'
    table4.rows[0].cells[2].text = 'Bevistema'
    for cell in table4.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    table4.rows[1].cells[0].text = 'V1'
    table4.rows[1].cells[1].text = 'Abu Mohammad (penningväxlare)'
    table4.rows[1].cells[2].text = 'Hebas hawala-verksamhet, transaktionsflöden, dollarkurser'
    table4.rows[2].cells[0].text = 'V2'
    table4.rows[2].cells[1].text = 'Mohammad Sami Alsharef (partsutsaga)'
    table4.rows[2].cells[2].text = 'Samtliga omständigheter i målet'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Angående kärandens vittnen: ').bold = True
    p.add_run('Heba har uppgett 2 vittnen vars identitet ännu inte avslöjats. Vi förbehåller oss rätten att bemöta deras vittnesmål.')
    
    doc.add_paragraph()
    doc.add_paragraph('Mohammad och Joumana förbehåller sig rätten att åberopa ytterligare bevisning före fristen 2026-03-20.')
    
    # SÄRSKILT OM JOUMANA
    doc.add_heading('SÄRSKILT OM JOUMANA ALNABLSI', level=1)
    
    doc.add_paragraph('Joumana har minimal personlig inblandning i denna sak. De 30 000 kr som hon påstås ha fått:')
    doc.add_paragraph('Var del av familjens hawala-verksamhet', style='List Number')
    doc.add_paragraph('Återbetalades via Swish', style='List Number')
    doc.add_paragraph('Hanterades av Mohammad', style='List Number')
    
    doc.add_paragraph('Joumana har utfärdat fullmakt till Mohammad att företräda henne i detta mål.')
    doc.add_paragraph('Vi hemställer att rätten överväger att avvisa kravet mot Joumana då hon inte var part i något avtal med Heba.')
    
    # GENKÄROMÅL
    doc.add_heading('GENKÄROMÅL – 400 000 KR', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Mohammad framställer genkäromål mot Heba Alhussien på totalt 400 000 kr.').bold = True
    
    doc.add_paragraph()
    doc.add_heading('A. Fordran: 35 000 kr (Hebas skulderkännande)', level=2)
    doc.add_paragraph('Den 29 december 2024 skrev Heba i WhatsApp:')
    
    add_arabic_with_translation(
        doc,
        'عليي سداد دين لابو سمير 35 الف كرون',
        '"Jag är skyldig Abu Sami [Mohammad] 35 000 kr"'
    )
    
    doc.add_paragraph('Meddelandet utgör ett ensidigt skulderkännande. Heba har inte betalat.')
    
    doc.add_heading('B. Kränkningsersättning: 150 000 kr (falsk våldtäktsanmälan)', level=2)
    doc.add_paragraph('Heba anmälde Mohammad för våldtäkt – påstod upprepade övergrepp under ~2 år. Polisen lade ned inom ~10 dagar p.g.a. total brist på bevis.')
    doc.add_paragraph()
    doc.add_paragraph('Rättslig grund:').runs[0].bold = True
    doc.add_paragraph('Falsk tillvitelse enl. 15 kap 7 § BrB – falskeligen tillvita brott inför myndighet', style='List Bullet')
    doc.add_paragraph('Kränkningsersättning enl. 2 kap 3 § SkL – allvarlig kränkning genom brott', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Beloppet 150 000 kr motiveras av:').runs[0].bold = True
    doc.add_paragraph('Våldtäkt är bland de allvarligaste brotten – falsk anklagelse utgör extrem kränkning med risk för frihetsberövande (min. 2 års fängelse)', style='List Bullet')
    doc.add_paragraph('Allvarlig skada på heder och rykte i både svenskt och arabiskt samhälle', style='List Bullet')
    doc.add_paragraph('Långvarigt psykiskt lidande – ångest, sömnproblem, påverkan på barn och familj', style='List Bullet')
    doc.add_paragraph('Systematisk förföljelse – del av hämndkampanj', style='List Bullet')
    doc.add_paragraph('Jfr Brottsoffermyndighetens praxis: 100 000–200 000 kr för falska anklagelser om allvarliga brott', style='List Bullet')
    
    doc.add_heading('C. Skadestånd: 100 000 kr (rättegångsmissbruk / falsk stämning)', level=2)
    doc.add_paragraph('Heba har medvetet lämnat in en stämning på 541 118 kr baserad på påståenden hon vet är falska. Pengarna var del av hawala-verksamhet – inga lån existerade.')
    doc.add_paragraph()
    doc.add_paragraph('Rättslig grund:').runs[0].bold = True
    doc.add_paragraph('Ren förmögenhetsskada vid brott enl. 2 kap 2 § SkL', style='List Bullet')
    doc.add_paragraph('Jfr osann partsutsaga enl. 15 kap 2 § BrB', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Beloppet 100 000 kr motiveras av:').runs[0].bold = True
    doc.add_paragraph('Fabricerat krav på 541 118 kr – systematiskt mönster (falsk våldtäktsanmälan, sedan falsk stämning)', style='List Bullet')
    doc.add_paragraph('Hundratals timmars arbete för att försvara sig mot grundlösa påståenden', style='List Bullet')
    doc.add_paragraph('Rättsväsendet används medvetet som hämndvapen', style='List Bullet')
    
    doc.add_heading('D. Ekonomisk skada: 75 000 kr', level=2)
    doc.add_paragraph('Direkt ekonomisk skada till följd av Hebas agerande:')
    doc.add_paragraph('Auktoriserade översättningar (arabiska → svenska): ~20 000 kr', style='List Bullet')
    doc.add_paragraph('Förlorad arbetsinkomst (samisjackets.com): ~30 000 kr', style='List Bullet')
    doc.add_paragraph('Resekostnader, porto, kopiering: ~10 000 kr', style='List Bullet')
    doc.add_paragraph('Framtida kostnader (huvudförhandling 18 maj 2026): ~15 000 kr', style='List Bullet')
    doc.add_paragraph('Rättslig grund: 2 kap 2 § SkL (ren förmögenhetsskada).')
    
    doc.add_heading('E. Sveda och värk: 40 000 kr (psykiskt lidande)', level=2)
    doc.add_paragraph('Mohammad har lidit allvarligt psykiskt lidande:')
    doc.add_paragraph('Ångest och sömnproblem sedan falska våldtäktsanmälan', style='List Bullet')
    doc.add_paragraph('Konstant stress från att vara instämd på 541 118 kr på falska grunder', style='List Bullet')
    doc.add_paragraph('Social skam, påverkan på föräldraskap och familjeliv', style='List Bullet')
    doc.add_paragraph('Rättslig grund: 5 kap 1 § SkL – personskada (sveda och värk).')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('TOTALT GENKÄROMÅL: 400 000 kr + ränta + rättegångskostnader').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Bevisning: ').bold = True
    p.add_run('B1 (skulderkännande), B5 (polisens nedläggningsbeslut), B6–B10 (Abu Mohammad, hawala-bevis), partsutsaga.')
    
    # SAMMANFATTNING
    doc.add_heading('SAMMANFATTNING', level=1)
    
    doc.add_paragraph('1. Inget lån existerar – det var en penningöverföringsverksamhet (hawala), bevisad genom Abu Mohammad-chatten', style='List Number')
    doc.add_paragraph('2. Heba är skyldig Mohammad 35 000 kr – enligt hennes eget skriftliga erkännande', style='List Number')
    doc.add_paragraph('3. Hebas trovärdighet är obefintlig – falsk våldtäktsanmälan avfärdad av polis inom 10 dagar', style='List Number')
    doc.add_paragraph('4. Joumana bör avvisas från målet – minimal inblandning, fullmakt till Mohammad', style='List Number')
    doc.add_paragraph('5. Abu Mohammad bekräftar Hebas hawala-verksamhet – 518 meddelanden med dollarväxlingar, USDT, fakturor', style='List Number')
    doc.add_paragraph('6. Genkäromål på 400 000 kr – skuld + kränkning + rättegångsmissbruk + ekonomisk skada + sveda och värk', style='List Number')
    
    # SLUTORD
    doc.add_heading('SLUTORD', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Stämningen är grundlös och motiverad av hämnd.').bold = True
    doc.add_paragraph()
    doc.add_paragraph('Heba har gjort en falsk våldtäktsanmälan (avfärdad av polisen), lämnat in en stämning grundad på lögner, dolt sin hawala-verksamhet och vägrat erkänna sin egen skuld på 35 000 kr.')
    doc.add_paragraph()
    doc.add_paragraph('Vi hemställer att rätten avslår käromålet i sin helhet och bifaller genkäromålet på 400 000 kr jämte ränta och rättegångskostnader.')
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph('Eskilstuna den 2026-02-16')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Mohammad Sami Alsharef').bold = True
    doc.add_paragraph('Även ombud för Joumana Alnablsi (enligt bifogad fullmakt)')
    doc.add_paragraph()
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('(Underskrift)')
    
    doc.add_paragraph()
    doc.add_paragraph('Bilagor:').runs[0].bold = True
    doc.add_paragraph('1. Fullmakt från Joumana Alnablsi')
    doc.add_paragraph('2. WhatsApp-utdrag med auktoriserad översättning')
    doc.add_paragraph('3. Abu Mohammad-chatt (utdrag)')
    doc.add_paragraph('4. Polisens nedläggningsbeslut')
    
    filepath = os.path.join(output_folder, '01_SVAROMAL_Mohammad_Sami_Alsharef.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# =============================================================================
# 2. SVAROMÅL - JOUMANA (UPDATED)
# =============================================================================
def create_svaromal_joumana():
    """Create Joumana's Svaromål - UPDATED"""
    doc = setup_doc()
    
    title = doc.add_heading('SVAROMÅL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Mål nr T 4438-25')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    doc.add_paragraph('Till:')
    doc.add_paragraph('Eskilstuna Tingsrätt\nRotel 2:04\n631 86 Eskilstuna')
    doc.add_paragraph()
    doc.add_paragraph('Datum: 2026-02-16')
    
    # PARTIES
    doc.add_heading('PARTER', level=1)
    
    doc.add_paragraph('Kärande:').runs[0].bold = True
    doc.add_paragraph('Heba Alhussien\nPersonnummer: 19890412-1244\nOmbud: Advokat Mayssam Baaken')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Svarande:').runs[0].bold = True
    doc.add_paragraph('Joumana Alnablsi\nPersonnummer: 19880102-5084\nNäshultagatan 12, 632 29 Eskilstuna\nE-post: jomana.alnablsi@hotmail.com\nTelefon: 0735165708')
    
    # OMBUD
    doc.add_heading('OMBUD', level=1)
    
    p = doc.add_paragraph('Jag har utfärdat ')
    p.add_run('fullmakt').bold = True
    p.add_run(' till ')
    p.add_run('Mohammad Sami Alsharef').bold = True
    p.add_run(' (19860217-5070) att företräda mig i detta mål och alla relaterade ärenden. Se bifogad fullmakt.')
    
    doc.add_paragraph('All kommunikation i detta mål ska ske genom mitt ombud Mohammad Sami Alsharef.')
    
    # INSTÄLLNING
    doc.add_heading('INSTÄLLNING', level=1)
    
    p = doc.add_paragraph('Jag ')
    p.add_run('bestrider').bold = True
    p.add_run(' käromålet i sin helhet. ')
    p.add_run('Ingen förlikning accepteras.').bold = True
    
    # KORT MOTIVERING
    doc.add_heading('KORT MOTIVERING', level=1)
    
    p = doc.add_paragraph('1. ')
    p.add_run('Jag har inte ingått något låneavtal').bold = True
    p.add_run(' med Heba Alhussien.')
    
    doc.add_paragraph()
    p = doc.add_paragraph('2. ')
    p.add_run('De 30 000 kr').bold = True
    p.add_run(' som påstås ha överförts till mig var del av en informell penningöverföringsverksamhet (hawala) som hanterades av Mohammad.')
    
    doc.add_paragraph()
    p = doc.add_paragraph('3. ')
    p.add_run('Beloppet återbetalades').bold = True
    p.add_run(' via Swish-överföringar, vilket framgår av WhatsApp-konversationer (B4).')
    
    doc.add_paragraph()
    p = doc.add_paragraph('4. ')
    p.add_run('Jag hade minimal personlig inblandning').bold = True
    p.add_run(' i de ekonomiska arrangemangen mellan Mohammad och Heba.')
    
    doc.add_paragraph()
    doc.add_paragraph('5. Jag ansluter mig i övrigt till det svaromål som inges av mitt ombud Mohammad Sami Alsharef.')
    
    # HEMSTÄLLAN
    doc.add_heading('HEMSTÄLLAN', level=1)
    
    doc.add_paragraph('Jag hemställer att:')
    p = doc.add_paragraph('1. ')
    p.add_run('Käromålet mot mig ogillas i sin helhet').bold = True
    
    p = doc.add_paragraph('2. ')
    p.add_run('Heba förpliktas ersätta mina rättegångskostnader').bold = True
    
    doc.add_paragraph('3. Alternativt att kravet mot mig avvisas då jag inte var part i något låneavtal')
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Eskilstuna den 2026-02-16')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Joumana Alnablsi').bold = True
    doc.add_paragraph()
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('(Underskrift)')
    
    doc.add_paragraph()
    doc.add_paragraph('Bilaga:').runs[0].bold = True
    doc.add_paragraph('Fullmakt till Mohammad Sami Alsharef')
    
    filepath = os.path.join(output_folder, '02_SVAROMAL_Joumana_Alnablsi.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# =============================================================================
# 3. FULLMAKT (UPDATED)
# =============================================================================
def create_fullmakt():
    """Create Fullmakt - UPDATED"""
    doc = setup_doc()
    
    title = doc.add_heading('FULLMAKT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Rättegångsfullmakt')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    # FULLMAKTSGIVARE
    doc.add_heading('FULLMAKTSGIVARE', level=1)
    
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Table Grid'
    table1.rows[0].cells[0].text = 'Namn:'
    table1.rows[0].cells[1].text = 'Joumana Alnablsi'
    table1.rows[1].cells[0].text = 'Personnummer:'
    table1.rows[1].cells[1].text = '19880102-5084'
    table1.rows[2].cells[0].text = 'Adress:'
    table1.rows[2].cells[1].text = 'Näshultagatan 12, 632 29 Eskilstuna'
    table1.rows[3].cells[0].text = 'Telefon:'
    table1.rows[3].cells[1].text = '0735165708'
    table1.rows[4].cells[0].text = 'E-post:'
    table1.rows[4].cells[1].text = 'jomana.alnablsi@hotmail.com'
    
    doc.add_paragraph()
    
    # FULLMAKTSTAGARE
    doc.add_heading('FULLMAKTSTAGARE (OMBUD)', level=1)
    
    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Table Grid'
    table2.rows[0].cells[0].text = 'Namn:'
    table2.rows[0].cells[1].text = 'Mohammad Sami Alsharef'
    table2.rows[1].cells[0].text = 'Personnummer:'
    table2.rows[1].cells[1].text = '19860217-5070'
    table2.rows[2].cells[0].text = 'Adress:'
    table2.rows[2].cells[1].text = 'Näshultagatan 12, 632 29 Eskilstuna'
    table2.rows[3].cells[0].text = 'Telefon:'
    table2.rows[3].cells[1].text = '0720147440'
    table2.rows[4].cells[0].text = 'E-post:'
    table2.rows[4].cells[1].text = 'info@samisjackets.com'
    table2.rows[5].cells[0].text = 'Relation:'
    table2.rows[5].cells[1].text = 'Före detta sambo'
    
    doc.add_paragraph()
    
    # FULLMAKTENS OMFATTNING
    doc.add_heading('FULLMAKTENS OMFATTNING', level=1)
    
    p = doc.add_paragraph('Jag, ')
    p.add_run('Joumana Alnablsi').bold = True
    p.add_run(', ger härmed ')
    p.add_run('Mohammad Sami Alsharef').bold = True
    p.add_run(' fullmakt att för min räkning:')
    
    doc.add_paragraph()
    doc.add_paragraph('Rättegångsfullmakt (Mål T 4438-25):').runs[0].bold = True
    
    doc.add_paragraph('Företräda mig som svarande i mål nr T 4438-25 vid Eskilstuna Tingsrätt', style='List Bullet')
    doc.add_paragraph('Avge svaromål, yttranden och övriga inlagor', style='List Bullet')
    doc.add_paragraph('Närvara vid muntlig förberedelse och huvudförhandling', style='List Bullet')
    doc.add_paragraph('Ta emot delgivning av domar och beslut', style='List Bullet')
    doc.add_paragraph('Överklaga domar och beslut', style='List Bullet')
    doc.add_paragraph('Vidta alla åtgärder som krävs för att tillvarata mina intressen i målet', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Generalfullmakt avseende detta ärende:').runs[0].bold = True
    
    doc.add_paragraph('Företräda mig inför alla myndigheter i anledning av detta mål', style='List Bullet')
    doc.add_paragraph('Ta emot och kvittera handlingar', style='List Bullet')
    doc.add_paragraph('Kommunicera med motparten och dennes ombud', style='List Bullet')
    doc.add_paragraph('Fatta beslut om processåtgärder', style='List Bullet')
    
    # SÄRSKILDA VILLKOR
    doc.add_heading('SÄRSKILDA VILLKOR', level=1)
    
    doc.add_paragraph('1. Denna fullmakt är obegränsad i tid och gäller tills den uttryckligen återkallas av mig.')
    doc.add_paragraph('2. Fullmaktstagaren har rätt att substituera fullmakten till juridiskt ombud om sådant anlitas.')
    doc.add_paragraph('3. Jag förbinder mig att godkänna de åtgärder som fullmaktstagaren vidtar inom ramen för denna fullmakt.')
    
    # SKÄL
    doc.add_heading('SKÄL TILL FULLMAKTEN', level=1)
    
    doc.add_paragraph('Mohammad Sami Alsharef har fullständig kunskap om omständigheterna; Joumana hade minimal inblandning i de ekonomiska arrangemangen med Heba. Gemensam företrädare är praktiskt.')
    
    # UNDERSKRIFTER
    doc.add_heading('UNDERSKRIFTER', level=1)
    
    doc.add_paragraph('Fullmaktsgivare:').runs[0].bold = True
    doc.add_paragraph()
    doc.add_paragraph('Ort och datum: Eskilstuna den ______________ 2026')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('Joumana Alnablsi (19880102-5084)')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Fullmaktstagare (bekräftelse av mottagande):').runs[0].bold = True
    doc.add_paragraph()
    doc.add_paragraph('Jag, Mohammad Sami Alsharef, bekräftar att jag har mottagit denna fullmakt och åtar mig att företräda Joumana Alnablsi i enlighet med dess villkor.')
    doc.add_paragraph()
    doc.add_paragraph('Ort och datum: Eskilstuna den ______________ 2026')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('Mohammad Sami Alsharef (19860217-5070)')
    
    filepath = os.path.join(output_folder, '03_FULLMAKT_Joumana_till_Mohammad.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# =============================================================================
# 4. BEVISUPPGIFT - WhatsApp (UPDATED with Abu Mohammad)
# =============================================================================
def create_whatsapp_evidence():
    """Create WhatsApp evidence document - UPDATED with Abu Mohammad evidence"""
    doc = setup_doc()
    
    title = doc.add_heading('BEVISUPPGIFT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('WhatsApp-konversationer som bevisning')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    
    subtitle2 = doc.add_paragraph('Mål T 4438-25 – Uppdaterad 2026-02-16')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('OBS: Alla arabiska texter kommer att förses med auktoriserad översättning innan fristen 2026-03-20.').bold = True
    p.runs[0].italic = True
    
    doc.add_paragraph()
    
    # === SECTION A: Heba-Mohammad chat ===
    doc.add_heading('A. KONVERSATION: HEBA – MOHAMMAD', level=1)
    
    table_info = doc.add_table(rows=4, cols=2)
    table_info.style = 'Table Grid'
    table_info.rows[0].cells[0].text = 'Källa:'
    table_info.rows[0].cells[1].text = 'WhatsApp-konversation mellan Mohammad och Heba'
    table_info.rows[1].cells[0].text = 'Period:'
    table_info.rows[1].cells[1].text = '2022 – februari 2025 (2 453 meddelanden)'
    table_info.rows[2].cells[0].text = 'Språk:'
    table_info.rows[2].cells[1].text = 'Arabiska (auktoriserad översättning bifogas)'
    table_info.rows[3].cells[0].text = 'Relevans:'
    table_info.rows[3].cells[1].text = 'Bevisar att Heba var skyldig Mohammad pengar, inte tvärtom'
    
    doc.add_paragraph()
    
    # BEVIS B1
    doc.add_heading('BEVIS B1: Hebas skulderkännande (29 december 2024)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Datum: 2024-12-29, kl 17:25:28').bold = True
    p.add_run('\nAvsändare: Heba Alhussien')
    
    add_arabic_with_translation(
        doc,
        'عليي سداد دين لابو سمير 35 الف كرون . لتاريخ اليوم بشهر ديسمبر وتم دفع الرينتا من ابو سمير كاش 1800 كرون ل هبه',
        '"Jag är skyldig Abu Sami [Mohammad] 35 000 kronor. Per dagens datum i december, och hyran betalades av Abu Sami kontant 1 800 kronor till Heba."'
    )
    
    p = doc.add_paragraph()
    p.add_run('BEVISVÄRDE: AVGÖRANDE').bold = True
    doc.add_paragraph('Heba erkänner skriftligen skuld till Mohammad 2 månader före stämningen. Grund för genkäromål.')
    
    # BEVIS B2
    doc.add_heading('BEVIS B2: Heba skyldig 60 000 kr (23 oktober 2024)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Datum: 2024-10-23, kl 11:58:41').bold = True
    p.add_run('\nAvsändare: Heba Alhussien')
    
    add_arabic_with_translation(
        doc,
        'كيف عليي ٦٠ الف ؟ شوووو الك عليي ٦٠ الف ؟ كيف',
        '"Hur är jag skyldig 60 000? VADÅ du säger jag är skyldig dig 60 000? Hur?"'
    )
    
    p = doc.add_paragraph()
    p.add_run('BEVISVÄRDE: STARKT').bold = True
    doc.add_paragraph('Heba ifrågasätter beloppet men inte att hon är skyldig – bekräftar Mohammads bokföring.')
    
    # BEVIS B3
    doc.add_heading('BEVIS B3: Hebas bokföring – $8 852 (juni 2023)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Datum: Juni 2023').bold = True
    p.add_run('\nAvsändare: Heba Alhussien')
    
    add_arabic_with_translation(
        doc,
        'حساب ابو سمير ٨٨٥٢ دولار',
        '"Abu Samis [Mohammads] konto: 8 852 dollar" (skuld till Mohammad)'
    )
    
    p = doc.add_paragraph()
    p.add_run('BEVISVÄRDE: STARKT').bold = True
    doc.add_paragraph('Hebas egen bokföring – typiskt för hawala-verksamhet, inte lån.')
    
    # BEVIS B4
    doc.add_heading('BEVIS B4: Återbetalning via Swish (januari 2023)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Datum: 2023-01-10, kl 14:16:02').bold = True
    p.add_run('\nAvsändare: Mohammad Sami Alsharef')
    
    add_arabic_with_translation(
        doc,
        'ربي يسلم ايديكي مبدائيا رجعتلك كل شي اخذتو منك سويش',
        '"Gud välsigne dina händer, jag har i princip lämnat tillbaka allt jag tog från dig via Swish"'
    )
    
    p = doc.add_paragraph()
    p.add_run('BEVISVÄRDE: STÖDJANDE').bold = True
    doc.add_paragraph('Pengar via Swish (inkl. 30 000 kr till Joumana) återbetalades.')
    
    # BEVIS – Bank Norwegian
    doc.add_heading('BEVIS: "DITT Bank Norwegian" (17 september 2024)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Datum: 2024-09-17, kl 09:22:57').bold = True
    p.add_run('\nAvsändare: Mohammad Sami Alsharef')
    
    add_arabic_with_translation(
        doc,
        'فوتي حوليلي من النوريجين تبعك فورا 35000 kr لحسابي',
        '"Gå in och överför till mig från DITT Norwegian [Bank Norwegian] omedelbart 35 000 kr till mitt konto"'
    )
    
    doc.add_paragraph('Visar att Bank Norwegian var Hebas eget konto ("تبعك" = "ditt").')
    
    doc.add_paragraph()
    
    # === SECTION B: Abu Mohammad chat ===
    doc.add_heading('B. KONVERSATION: HEBA – ABU MOHAMMAD (PENNINGVÄXLARE)', level=1)
    
    p = doc.add_paragraph()
    p.add_run('NY BEVISNING – åberopas före fristen 2026-03-20').bold = True
    
    doc.add_paragraph()
    
    table_abu = doc.add_table(rows=5, cols=2)
    table_abu.style = 'Table Grid'
    table_abu.rows[0].cells[0].text = 'Källa:'
    table_abu.rows[0].cells[1].text = 'WhatsApp-konversation: Heba ↔ Abu Mohammad (صراف = penningväxlare)'
    table_abu.rows[1].cells[0].text = 'Period:'
    table_abu.rows[1].cells[1].text = 'December 2023 – september 2024 (518 meddelanden)'
    table_abu.rows[2].cells[0].text = 'Innehåll:'
    table_abu.rows[2].cells[1].text = '60+ ljudfiler, 80+ foton, 15+ PDF-fakturor, 2 videor'
    table_abu.rows[3].cells[0].text = 'Språk:'
    table_abu.rows[3].cells[1].text = 'Arabiska (auktoriserad översättning bifogas)'
    table_abu.rows[4].cells[0].text = 'Relevans:'
    table_abu.rows[4].cells[1].text = 'Bevisar att Heba bedrev professionell hawala-verksamhet'
    
    doc.add_paragraph()
    
    # B6
    doc.add_heading('BEVIS B6: Dollarväxlingar i Damaskus', level=2)
    doc.add_paragraph('Flera meddelanden visar att Heba och Abu Mohammad genomförde valutaväxlingar:')
    
    add_arabic_with_translation(
        doc,
        'بدي بالشام ٤٠٠ دولار',
        '"Jag behöver 400 dollar i Damaskus"'
    )
    
    doc.add_paragraph('Detta visar att Heba använde sitt svenska konto för att betala i Sverige medan motsvarande belopp levererades i Damaskus – identiskt system som med Mohammad.')
    
    # B7
    doc.add_heading('BEVIS B7: Stora belopp i syriska pund', level=2)
    doc.add_paragraph('En enda bekräftelse visar:')
    
    add_arabic_with_translation(
        doc,
        'تم 11,520,000 ل.س',
        '"Klart: 11 520 000 syriska pund" (bekräftelse av genomförd transaktion)'
    )
    
    doc.add_paragraph('Ytterligare: "عندي شي ٣٠ الف" = "~30 000 kr" och "٢٥٤٧٠ كرون" = "25 470 kr" (med foto på betalningsbevis).')
    
    # B8
    doc.add_heading('BEVIS B8: USDT/kryptovaluta', level=2)
    doc.add_paragraph('Meddelanden visar att kryptovaluta (USDT) var en del av verksamheten:')
    
    add_arabic_with_translation(
        doc,
        'عندي usdt',
        '"Jag har USDT [kryptovaluta]"'
    )
    
    doc.add_paragraph('Bekräftar att det handskrivna papperet var relaterat till USDT-handel, inte låneavtal.')
    
    # B9
    doc.add_heading('BEVIS B9: Faktura- och Klarna-betalningar', level=2)
    doc.add_paragraph('Chatten innehåller 15+ PDF-fakturor från:')
    doc.add_paragraph('Klarna-fakturor', style='List Bullet')
    doc.add_paragraph('Nordea Ropo Capital-fakturor', style='List Bullet')
    doc.add_paragraph('Övriga svenska fakturor', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Samma mönster som med Mohammad: Heba betalade svenska fakturor, motprestationen skedde i Syrien.')
    
    # B10
    doc.add_heading('BEVIS B10: Ljudfiler och fotografier', level=2)
    doc.add_paragraph('Chatten innehåller:')
    doc.add_paragraph('60+ ljudmeddelanden (PTT) med transaktionsinstruktioner', style='List Bullet')
    doc.add_paragraph('80+ fotografier med betalningsbevis och kvitton', style='List Bullet')
    doc.add_paragraph('2 videor', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Kan spelas upp i rätten som bevisning om Hebas professionella verksamhet.')
    
    # SLUTSATS
    doc.add_heading('SLUTSATS', level=1)
    
    doc.add_paragraph('WhatsApp-konversationerna visar med all tydlighet att:')
    doc.add_paragraph('1. Heba erkände skriftligen att hon var skyldig Mohammad 35 000 kr (B1)', style='List Number')
    doc.add_paragraph('2. Heba förde bokföring över skulder – typiskt för hawala (B3)', style='List Number')
    doc.add_paragraph('3. Bank Norwegian var Hebas eget konto, inte ett lån', style='List Number')
    doc.add_paragraph('4. Pengar via Swish återbetalades (B4)', style='List Number')
    doc.add_paragraph('5. Heba bedrev identisk hawala-verksamhet med Abu Mohammad (B6–B10)', style='List Number')
    doc.add_paragraph('6. Hebas påståenden i stämningen är lögner', style='List Number')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Heba var inte en långivare – hon var en professionell penningväxlare som nu ljuger inför rätten.').bold = True
    
    filepath = os.path.join(output_folder, '04_BEVISUPPGIFT_WhatsApp_FIXED.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# =============================================================================
# 5. NYCKELBEVIS - Hebas skulderkännande (UPDATED)
# =============================================================================
def create_key_evidence():
    """Create key evidence document - the smoking gun"""
    doc = setup_doc()
    
    title = doc.add_heading('NYCKELBEVIS B1', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Hebas skulderkännande – 35 000 kr')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(16)
    
    p = doc.add_paragraph('Grund för genkäromål')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Mål:'
    table.rows[0].cells[1].text = 'T 4438-25, Eskilstuna Tingsrätt'
    table.rows[1].cells[0].text = 'Datum för meddelande:'
    table.rows[1].cells[1].text = '29 december 2024, kl 17:25:28'
    table.rows[2].cells[0].text = 'Avsändare:'
    table.rows[2].cells[1].text = 'Heba Alhussien (käranden)'
    table.rows[3].cells[0].text = 'Mottagare:'
    table.rows[3].cells[1].text = 'Mohammad Sami Alsharef (svaranden)'
    table.rows[4].cells[0].text = 'Kanal:'
    table.rows[4].cells[1].text = 'WhatsApp'
    
    doc.add_paragraph()
    
    # THE MESSAGE
    doc.add_heading('ORIGINALMEDDELANDE (ARABISKA)', level=1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    arabic_run = p.add_run('عليي سداد دين لابو سمير 35 الف كرون')
    set_arabic_font(arabic_run, 'Arial', 18)
    arabic_run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    arabic_run2 = p.add_run('لتاريخ اليوم بشهر ديسمبر وتم دفع الرينتا من ابو سمير كاش 1800 كرون ل هبه')
    set_arabic_font(arabic_run2, 'Arial', 14)
    
    doc.add_paragraph()
    
    doc.add_heading('SVENSK ÖVERSÄTTNING', level=1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('"Jag är skyldig Abu Sami [Mohammad] 35 000 kronor."')
    run.font.size = Pt(18)
    run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('"Per dagens datum i december, och hyran betalades av Abu Sami kontant 1 800 kronor till Heba."')
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('OBS: Auktoriserad översättning kommer att bifogas före fristen 2026-03-20.').italic = True
    
    doc.add_paragraph()
    
    # Explanation
    doc.add_heading('FÖRKLARING AV TERMER', level=1)
    
    table2 = doc.add_table(rows=4, cols=2)
    table2.style = 'Table Grid'
    
    table2.rows[0].cells[0].text = 'Arabiska'
    table2.rows[0].cells[1].text = 'Betydelse'
    for cell in table2.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    
    arabic_cell = table2.rows[1].cells[0]
    arabic_para = arabic_cell.paragraphs[0]
    arabic_run = arabic_para.add_run('ابو سمير')
    set_arabic_font(arabic_run, 'Arial', 11)
    table2.rows[1].cells[1].text = '"Abu Sami" = Mohammad Sami Alsharef (smeknamn baserat på äldsta sonens namn)'
    
    arabic_cell2 = table2.rows[2].cells[0]
    arabic_para2 = arabic_cell2.paragraphs[0]
    arabic_run2 = arabic_para2.add_run('عليي سداد دين')
    set_arabic_font(arabic_run2, 'Arial', 11)
    table2.rows[2].cells[1].text = '"Jag är skyldig att betala en skuld" – ensidigt skulderkännande'
    
    arabic_cell3 = table2.rows[3].cells[0]
    arabic_para3 = arabic_cell3.paragraphs[0]
    arabic_run3 = arabic_para3.add_run('الرينتا')
    set_arabic_font(arabic_run3, 'Arial', 11)
    table2.rows[3].cells[1].text = '"Hyran" (från svenska "ränta/hyra")'
    
    doc.add_paragraph()
    
    # Significance
    doc.add_heading('BETYDELSE FÖR MÅLET', level=1)
    
    doc.add_paragraph('Detta meddelande skrevs av Heba Alhussien den 29 december 2024 – drygt 2 månader innan hon lämnade in stämningsansökan.')
    doc.add_paragraph()
    doc.add_paragraph('Heba erkänner uttryckligen att:')
    doc.add_paragraph('Hon ÄR SKYLDIG Mohammad ("Abu Sami") pengar', style='List Number')
    doc.add_paragraph('Skulden uppgår till 35 000 kr', style='List Number')
    doc.add_paragraph('Mohammad betalade hennes hyra (1 800 kr kontant)', style='List Number')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Detta motsäger helt Hebas påstående i stämningen att Mohammad är skyldig henne pengar.').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Mohammad yrkar i genkäromål återbetalning av dessa 35 000 kr.').bold = True
    
    doc.add_paragraph()
    
    # Certification
    doc.add_paragraph('Jag intygar att ovanstående är en korrekt återgivning av WhatsApp-meddelandet:')
    doc.add_paragraph()
    doc.add_paragraph('Eskilstuna den 2026-02-16')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('Mohammad Sami Alsharef')
    
    filepath = os.path.join(output_folder, '05_NYCKELBEVIS_Hebas_Skulderkannande_FIXED.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# =============================================================================
# 6. YTTRANDE - NY BEVISNING (NEW DOCUMENT)
# =============================================================================
def create_yttrande_ny_bevisning():
    """Create formal court submission about new evidence"""
    doc = setup_doc()
    
    title = doc.add_heading('YTTRANDE AVSEENDE NY BEVISNING', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Mål nr T 4438-25')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    doc.add_paragraph('Till:')
    doc.add_paragraph('Eskilstuna Tingsrätt\nRotel 2:04\n631 86 Eskilstuna')
    doc.add_paragraph()
    doc.add_paragraph('Datum: 2026-02-16')
    doc.add_paragraph('Sista dag för ny bevisning: 2026-03-20')
    
    doc.add_paragraph()
    
    # BAKGRUND
    doc.add_heading('BAKGRUND', level=1)
    doc.add_paragraph('Vid muntlig förberedelse den 16 februari 2026 redovisades att:')
    doc.add_paragraph('Heba Alhussien har åberopat 2 vittnen vars identitet ännu inte avslöjats', style='List Bullet')
    doc.add_paragraph('Parterna har frist till den 20 mars 2026 att inkomma med ny bevisning', style='List Bullet')
    doc.add_paragraph('All arabisk bevisning ska förses med auktoriserad översättning', style='List Bullet')
    
    # NY BEVISNING
    doc.add_heading('NY BEVISNING SOM ÅBEROPAS', level=1)
    
    doc.add_heading('1. Abu Mohammad-chatten (B6–B10)', level=2)
    doc.add_paragraph('WhatsApp-konversation mellan Heba och Abu Mohammad (penningväxlare/صراف), dec 2023–sep 2024.')
    doc.add_paragraph()
    doc.add_paragraph('Bevistema:').runs[0].bold = True
    doc.add_paragraph('Heba bedrev systematisk hawala med professionella penningväxlare', style='List Bullet')
    doc.add_paragraph('Påstådda "lån" var del av denna verksamhet', style='List Bullet')
    doc.add_paragraph('Identiskt transaktionsmönster (fakturor, Swish, dollarkurser) som med Mohammad', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Chatten innehåller:')
    doc.add_paragraph('518 meddelanden, 60+ ljud, 80+ foton, 15+ PDF-fakturor', style='List Bullet')
    doc.add_paragraph('Dollarväxlingar, 11,5 milj. SYP, USDT/krypto', style='List Bullet')
    
    # VITTNE
    doc.add_heading('2. Abu Mohammad som vittne', level=2)
    doc.add_paragraph('Vi åberopar Abu Mohammad som vittne.')
    doc.add_paragraph()
    doc.add_paragraph('Bevistema:').runs[0].bold = True
    doc.add_paragraph('Hebas hawala-verksamhet och transaktionsflöden', style='List Bullet')
    doc.add_paragraph('Dollarkurser, betalningsflöden och internationella transaktioner', style='List Bullet')
    
    # GENKÄROMÅL
    doc.add_heading('3. Genkäromål – 400 000 kr', level=2)
    doc.add_paragraph('Mohammad framställer genkäromål mot Heba Alhussien på totalt 400 000 kr:')
    doc.add_paragraph('35 000 kr – Hebas skuld (skulderkännande 2024-12-29)', style='List Bullet')
    doc.add_paragraph('150 000 kr – Kränkningsersättning för falsk våldtäktsanmälan (2 kap 3 § SkL, BrB 15:7)', style='List Bullet')
    doc.add_paragraph('100 000 kr – Skadestånd för rättegångsmissbruk (2 kap 2 § SkL, jfr BrB 15:2)', style='List Bullet')
    doc.add_paragraph('75 000 kr – Ekonomisk skada (översättningskostnader, förlorad inkomst, resekostnader m.m.)', style='List Bullet')
    doc.add_paragraph('40 000 kr – Sveda och värk / psykiskt lidande (5 kap 1 § SkL)', style='List Bullet')
    doc.add_paragraph('Se separat genkäromålsinlaga (08_GENKAROMAL_400000kr.docx) för fullständiga grunder och rättsliga hänvisningar.')
    
    # ANGÅENDE KÄRANDENS VITTNEN
    doc.add_heading('ANGÅENDE KÄRANDENS VITTNEN', level=1)
    doc.add_paragraph('Heba Alhussien har uppgett 2 vittnen. Vi begär att:')
    doc.add_paragraph('1. Vittnenas identitet redovisas snarast', style='List Number')
    doc.add_paragraph('2. Bevistema för varje vittne anges', style='List Number')
    doc.add_paragraph('3. Vi ges möjlighet att ställa frågor till vittnena vid huvudförhandlingen', style='List Number')
    
    # TRANSLATION NOTE
    doc.add_heading('AUKTORISERAD ÖVERSÄTTNING', level=1)
    doc.add_paragraph('I enlighet med rättens anvisningar förses all arabisk bevisning med auktoriserad svensk översättning före fristen 2026-03-20.')
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph('Eskilstuna den 2026-02-16')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Mohammad Sami Alsharef').bold = True
    doc.add_paragraph('Svarande 1, även ombud för Joumana Alnablsi')
    doc.add_paragraph()
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('(Underskrift)')
    
    filepath = os.path.join(output_folder, '07_YTTRANDE_NY_BEVISNING.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# =============================================================================
# 7. GENKÄROMÅL (NEW DOCUMENT)
# =============================================================================
def create_genkaromal():
    """Create formal counterclaim document - 150 000 kr"""
    doc = setup_doc()
    
    title = doc.add_heading('GENKÄROMÅL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Mål nr T 4438-25')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Totalt krav: 400 000 kr + ränta + rättegångskostnader')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    doc.add_paragraph()
    
    doc.add_paragraph('Till:')
    doc.add_paragraph('Eskilstuna Tingsrätt\nRotel 2:04\n631 86 Eskilstuna')
    doc.add_paragraph()
    doc.add_paragraph('Datum: 2026-02-16')
    
    # PARTER
    doc.add_heading('PARTER I GENKÄROMÅLET', level=1)
    
    doc.add_paragraph('Genkärande:').runs[0].bold = True
    doc.add_paragraph('Mohammad Sami Alsharef\nPersonnummer: 19860217-5070\nNäshultagatan 12, 632 29 Eskilstuna\nE-post: info@samisjackets.com\nTelefon: 0720147440')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Gensvarande:').runs[0].bold = True
    doc.add_paragraph('Heba Alhussien\nPersonnummer: 19890412-1244\nOmbud: Advokat Mayssam Baaken')
    
    # YRKANDE
    doc.add_heading('YRKANDE', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Mohammad Sami Alsharef yrkar att Heba Alhussien förpliktas att betala sammanlagt 400 000 kr till Mohammad Sami Alsharef, fördelat enligt följande:').bold = True
    doc.add_paragraph()
    
    # Yrkande table
    table_y = doc.add_table(rows=7, cols=3)
    table_y.style = 'Table Grid'
    table_y.rows[0].cells[0].text = 'Post'
    table_y.rows[0].cells[1].text = 'Belopp'
    table_y.rows[0].cells[2].text = 'Rättslig grund'
    for cell in table_y.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    
    y_data = [
        ('A. Fordran (Hebas skuld)', '35 000 kr', 'Ensidigt skulderkännande 2024-12-29'),
        ('B. Kränkningsersättning', '150 000 kr', '2 kap 3 § SkL via BrB 15:7'),
        ('C. Skadestånd rättegångsmissbruk', '100 000 kr', '2 kap 2 § SkL, jfr BrB 15:2'),
        ('D. Ekonomisk skada', '75 000 kr', '2 kap 2 § SkL'),
        ('E. Sveda och värk', '40 000 kr', '5 kap 1 § SkL'),
        ('TOTALT', '400 000 kr', '+ ränta + rättegångskostnader'),
    ]
    
    for i, (post, belopp, grund) in enumerate(y_data, 1):
        table_y.rows[i].cells[0].text = post
        table_y.rows[i].cells[1].text = belopp
        table_y.rows[i].cells[2].text = grund
        if i == 6:
            for cell in table_y.rows[i].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Jämte ränta enligt 6 § räntelagen (1975:635) från respektive skadetillfälle till dess betalning sker.')
    doc.add_paragraph('Heba förpliktas att ersätta samtliga rättegångskostnader.')
    
    # GRUNDER
    doc.add_heading('GRUNDER', level=1)
    
    # A. Fordran
    doc.add_heading('A. Fordran: 35 000 kr', level=2)
    
    doc.add_paragraph('1. Den 29 december 2024 skrev Heba Alhussien följande i ett WhatsApp-meddelande till Mohammad:')
    doc.add_paragraph()
    
    add_arabic_with_translation(
        doc,
        'عليي سداد دين لابو سمير 35 الف كرون . لتاريخ اليوم بشهر ديسمبر وتم دفع الرينتا من ابو سمير كاش 1800 كرون ل هبه',
        '"Jag är skyldig Abu Sami [Mohammad] 35 000 kr. Per dagens datum i december, och hyran betalades av Abu Sami kontant 1 800 kr till Heba."'
    )
    
    doc.add_paragraph('2. Meddelandet utgör ett ensidigt skulderkännande. Heba erkänner att hon är skyldig Mohammad 35 000 kr.')
    doc.add_paragraph('3. Meddelandet skrevs den 29 december 2024 – drygt 2 månader innan Heba lämnade in sin stämningsansökan.')
    doc.add_paragraph('4. Heba har inte betalat skulden.')
    
    # B. Kränkningsersättning
    doc.add_heading('B. Kränkningsersättning: 150 000 kr (falsk våldtäktsanmälan)', level=2)
    
    doc.add_paragraph('1. Heba anmälde Mohammad för våldtäkt – påstod upprepade övergrepp under ~2 år.')
    doc.add_paragraph()
    doc.add_paragraph('2. Polisen lade ned inom ~10 dagar p.g.a. total brist på bevis.')
    doc.add_paragraph()
    doc.add_paragraph('3. Falsk tillvitelse enl. 15 kap 7 § BrB. Kränkningsersättning enl. 2 kap 3 § SkL.')
    doc.add_paragraph()
    doc.add_paragraph('4. Mohammad har lidit allvarlig kränkning:')
    doc.add_paragraph('Extremt allvarlig anklagelse med risk för frihetsberövande (min. 2 års fängelse)', style='List Bullet')
    doc.add_paragraph('Allvarlig skada på heder och rykte i både svenskt och arabiskt samhälle', style='List Bullet')
    doc.add_paragraph('Långvarigt psykiskt lidande, påverkan på barn och familj', style='List Bullet')
    doc.add_paragraph('Del av systematisk hämndkampanj', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('5. 150 000 kr är skäligt. Jfr Brottsoffermyndighetens praxis: 100 000–200 000 kr för falska anklagelser om allvarliga brott.')
    
    # C. Rättegångsmissbruk
    doc.add_heading('C. Skadestånd: 100 000 kr (rättegångsmissbruk / falsk stämning)', level=2)
    
    doc.add_paragraph('1. Heba har medvetet lämnat in en stämning på 541 118 kr baserad på påståenden hon vet är falska.')
    doc.add_paragraph()
    doc.add_paragraph('2. Inga lån existerade – pengarna var del av hawala-verksamhet (bevisat genom Abu Mohammad-chatten).')
    doc.add_paragraph()
    doc.add_paragraph('3. Systematiskt mönster:')
    doc.add_paragraph('Först: Falsk våldtäktsanmälan (nedlagd)', style='List Bullet')
    doc.add_paragraph('Sedan: Falsk stämning på 541 118 kr', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('4. Rättslig grund: Ren förmögenhetsskada vid brott enl. 2 kap 2 § SkL. Jfr osann partsutsaga enl. 15 kap 2 § BrB.')
    doc.add_paragraph()
    doc.add_paragraph('5. Beloppet 100 000 kr motiveras av fabricerat krav, systematiskt missbruk och hundratals timmars tvingat försvar.')
    
    # D. Ekonomisk skada
    doc.add_heading('D. Ekonomisk skada: 75 000 kr', level=2)
    
    doc.add_paragraph('Mohammad har lidit direkt ekonomisk skada till följd av Hebas agerande:')
    doc.add_paragraph()
    
    table_ek = doc.add_table(rows=7, cols=2)
    table_ek.style = 'Table Grid'
    table_ek.rows[0].cells[0].text = 'Post'
    table_ek.rows[0].cells[1].text = 'Uppskattad kostnad'
    for cell in table_ek.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    table_ek.rows[1].cells[0].text = 'Auktoriserade översättningar (arabiska → svenska)'
    table_ek.rows[1].cells[1].text = '15 000 – 20 000 kr'
    table_ek.rows[2].cells[0].text = 'Förlorad arbetsinkomst (eget företag: samisjackets.com)'
    table_ek.rows[2].cells[1].text = '25 000 – 30 000 kr'
    table_ek.rows[3].cells[0].text = 'Resekostnader till domstol, möten, myndigheter'
    table_ek.rows[3].cells[1].text = '3 000 – 5 000 kr'
    table_ek.rows[4].cells[0].text = 'Porto, kopiering, administration'
    table_ek.rows[4].cells[1].text = '2 000 – 3 000 kr'
    table_ek.rows[5].cells[0].text = 'Framtida kostnader (huvudförhandling 18 maj 2026)'
    table_ek.rows[5].cells[1].text = '10 000 – 15 000 kr'
    table_ek.rows[6].cells[0].text = 'TOTALT (yrkat)'
    table_ek.rows[6].cells[1].text = '75 000 kr'
    for cell in table_ek.rows[6].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Rättslig grund: 2 kap 2 § skadeståndslagen – ren förmögenhetsskada. Kvitton och verifikationer kan presenteras.')
    
    # E. Sveda och värk
    doc.add_heading('E. Sveda och värk: 40 000 kr (psykiskt lidande)', level=2)
    
    doc.add_paragraph('Mohammad har lidit allvarligt psykiskt lidande till följd av Hebas systematiska förföljelse:')
    doc.add_paragraph('Ångest och sömnproblem sedan falska våldtäktsanmälan', style='List Bullet')
    doc.add_paragraph('Konstant stress från att vara instämd på 541 118 kr på falska grunder', style='List Bullet')
    doc.add_paragraph('Social skam, påverkan på föräldraskap och familjeliv', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Rättslig grund: 5 kap 1 § SkL – ersättning för personskada (sveda och värk, psykiskt lidande av övergående natur).')
    
    # HEBAS MÖNSTER
    doc.add_heading('HEBAS MÖNSTER AV MISSBRUK AV RÄTTSVÄSENDET', level=1)
    
    doc.add_paragraph('Det är viktigt att rätten ser helhetsbilden av Hebas agerande:')
    doc.add_paragraph()
    
    table_m = doc.add_table(rows=4, cols=3)
    table_m.style = 'Table Grid'
    table_m.rows[0].cells[0].text = 'Nr'
    table_m.rows[0].cells[1].text = 'Hebas åtgärd'
    table_m.rows[0].cells[2].text = 'Resultat'
    for cell in table_m.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    table_m.rows[1].cells[0].text = '1'
    table_m.rows[1].cells[1].text = 'Falsk våldtäktsanmälan till polisen'
    table_m.rows[1].cells[2].text = 'Nedlagd inom ~10 dagar – total brist på bevis'
    table_m.rows[2].cells[0].text = '2'
    table_m.rows[2].cells[1].text = 'Falsk stämning på 541 118 kr'
    table_m.rows[2].cells[2].text = 'Pågår – T 4438-25'
    table_m.rows[3].cells[0].text = '3'
    table_m.rows[3].cells[1].text = 'Allt sker EFTER att Mohammad blockerade Heba'
    table_m.rows[3].cells[2].text = 'Visar hämndmotiv'
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Heba använder rättsväsendet systematiskt som vapen. Detta mönster ska få konsekvenser.').bold = True
    
    # BEVISNING
    doc.add_heading('BEVISNING', level=1)
    
    doc.add_paragraph('Skriftlig bevisning:').runs[0].bold = True
    
    table_b = doc.add_table(rows=6, cols=3)
    table_b.style = 'Table Grid'
    table_b.rows[0].cells[0].text = 'Nr'
    table_b.rows[0].cells[1].text = 'Bevis'
    table_b.rows[0].cells[2].text = 'Bevisar'
    for cell in table_b.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    
    b_data = [
        ('B1', 'WhatsApp 2024-12-29 (auktoriserad översättning)', 'Hebas skulderkännande: 35 000 kr'),
        ('B5', 'Polisens nedläggningsbeslut', 'Falsk våldtäktsanmälan → kränkningsersättning'),
        ('B6–B10', 'Abu Mohammad-chatten (518 meddelanden)', 'Hawala-verksamhet → stämningen är falsk'),
        ('Kvitton', 'Översättningskostnader, inkomstförlust, resekostnader m.m.', 'Ekonomisk skada: 75 000 kr'),
        ('V2', 'Mohammad Sami Alsharef (partsutsaga)', 'Samtliga omständigheter'),
    ]
    
    for i, (nr, bevis, bevisar) in enumerate(b_data, 1):
        table_b.rows[i].cells[0].text = nr
        table_b.rows[i].cells[1].text = bevis
        table_b.rows[i].cells[2].text = bevisar
    
    # RÄTTSLIGA GRUNDER SAMMANFATTNING
    doc.add_heading('RÄTTSLIGA GRUNDER – SAMMANFATTNING', level=1)
    
    doc.add_paragraph('Skadeståndslagen (1972:207):').runs[0].bold = True
    doc.add_paragraph('2 kap 2 § – Ren förmögenhetsskada vid brott', style='List Bullet')
    doc.add_paragraph('2 kap 3 § – Kränkningsersättning vid allvarlig kränkning', style='List Bullet')
    doc.add_paragraph('5 kap 1 § – Sveda och värk', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Brottsbalken:').runs[0].bold = True
    doc.add_paragraph('15 kap 7 § – Falsk tillvitelse', style='List Bullet')
    doc.add_paragraph('15 kap 2 § – Osann partsutsaga', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Övrigt:').runs[0].bold = True
    doc.add_paragraph('6 § räntelagen (1975:635) – ränta från skadetillfälle', style='List Bullet')
    doc.add_paragraph('18 kap RB – rättegångskostnader', style='List Bullet')
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Eskilstuna den 2026-02-16')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Mohammad Sami Alsharef').bold = True
    doc.add_paragraph()
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('(Underskrift)')
    
    filepath = os.path.join(output_folder, '08_GENKAROMAL_400000kr.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


# =============================================================================
# RUN ALL
# =============================================================================
if __name__ == "__main__":
    print("=" * 60)
    print("CREATING WORD DOCUMENTS - UPDATED AFTER COURT HEARING")
    print("Case T 4438-25 - Eskilstuna Tingsrätt")
    print("Updated: 2026-02-16")
    print("Next hearing: 2026-05-18")
    print("Evidence deadline: 2026-03-20")
    print("=" * 60)
    print()
    
    create_svaromal_mohammad()
    create_svaromal_joumana()
    create_fullmakt()
    create_whatsapp_evidence()
    create_key_evidence()
    create_yttrande_ny_bevisning()
    create_genkaromal()
    
    print()
    print("=" * 60)
    print(f"ALL DOCUMENTS CREATED IN: {output_folder}")
    print("=" * 60)
    print()
    print("DOCUMENTS:")
    print("1. 01_SVAROMAL_Mohammad  - Updated with Abu Mohammad + genkäromål")
    print("2. 02_SVAROMAL_Joumana   - Updated dates, no settlement")
    print("3. 03_FULLMAKT           - Updated dates")
    print("4. 04_BEVISUPPGIFT       - Updated with B6-B10 Abu Mohammad evidence")
    print("5. 05_NYCKELBEVIS        - Updated with genkäromål reference")
    print("6. 07_YTTRANDE           - NEW: Formal submission about new evidence")
    print("7. 08_GENKAROMAL         - Formal counterclaim 400 000 kr")
    print()
    print("NEXT STEPS:")
    print("- Get authorized translations before 2026-03-20")
    print("- Contact Abu Mohammad for witness statement")
    print("- File all documents with the court")
